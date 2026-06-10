"""
test_cbr_docx.py — Harness de auditoría en frío del filtro CBR (Audit 11).
Reutiliza segmentación real del DOCX y evaluar_antiguedad_cbr de filtro_cbr.py.
Instrumenta TODOS los candidatos de año para que el abogado audite sin API.

Uso: python test_cbr_docx.py "<ruta_docx>"
Vuelca: auditoria_cbr_YYYY-MM-DD.csv en el directorio actual.
"""

import csv
import logging
import os
import re
import sys
from datetime import datetime

# Agregar el directorio del script al path para importar módulos del proyecto
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document

# -- Reutilizar segmentación real --
from modulo1_parser import (
    separar_bloques,
    extraer_rol,
    # regex de segmentación del DOCX (module-level, importables)
    _RE_FECHA_ENCABEZADO,
    _RE_TITULO_SEMANA,
)

# -- Reutilizar evaluación CBR real + constantes de regex --
from filtro_cbr import (
    evaluar_antiguedad_cbr,
    _CBR_ANCLAS,
    _CBR_ANIO_TAG,
    _CBR_VIGENTE,
    _CBR_ANIO_MIN,
    _CBR_ANIO_MAX,
    _CBR_VENTANA,
    _CBR_ANIO_CORTE,
)

# Réplica de _RE_TIENE_ROL de v2_experimental/modulo1_v2.py línea 83
# (no se importa el módulo v2 para evitar efectos de import en un
#  harness sin API). Mantener sincronizado si v2 lo cambia.
_RE_TIENE_ROL = re.compile(r'[CcVvAa]\s*[-–—]\s*\d+\s*[-–—]\s*\d{4}')

# ---------------------------------------------------------------------------
# Instrumentación: captura TODOS los candidatos de año (la función original
# solo devuelve el elegido; el harness necesita transparencia total).
# NO modifica filtro_cbr.py: replica la misma lógica para exponer los datos.
# ---------------------------------------------------------------------------

def _instrumentar_candidatos(bloque_texto: str) -> dict:
    """
    Ejecuta la misma lógica de candidatos que evaluar_antiguedad_cbr pero
    devuelve un dict con LA LISTA COMPLETA de años que matchearon, cuáles
    entraron como candidatos (cerca de ancla), y la marca vigente si existe.
    """
    texto = re.sub(r"\s+", " ", (bloque_texto or "").lower())

    # 1. Anclas CBR
    anclas = [(m.start(), m.group()) for m in _CBR_ANCLAS.finditer(texto)]
    posiciones_anclas = [ap for ap, _ in anclas]

    # 2. Todos los años que matchean _CBR_ANIO_TAG (dentro de rango)
    todos_los_anios = []
    for m in _CBR_ANIO_TAG.finditer(texto):
        a = int(m.group(1))
        if _CBR_ANIO_MIN <= a <= _CBR_ANIO_MAX:
            todos_los_anios.append({
                "anio": a,
                "pos": m.start(),
                "texto_matc": m.group(0),
            })

    # 3. De esos, cuáles están a ≤150 chars de una ancla → candidatos
    candidatos = []
    for item in todos_los_anios:
        a = item["anio"]
        pos = item["pos"]
        if not posiciones_anclas:
            continue
        dist_min = min(abs(pos - ap) for ap in posiciones_anclas)
        if dist_min <= _CBR_VENTANA:
            item_copy = dict(item)
            item_copy["distancia_ancla"] = dist_min
            candidatos.append(item_copy)

    # 4. Marca vigente
    vigente_pos = None
    m_vig = _CBR_VIGENTE.search(texto)
    if m_vig:
        vigente_pos = m_vig.start()

    # 5. Calcular dist_a_vigente para cada candidato
    for c in candidatos:
        d_vig = abs(c["pos"] - vigente_pos) if vigente_pos is not None else 9999
        c["distancia_vigente"] = d_vig

    vigente = {"pos": vigente_pos, "texto": m_vig.group()} if m_vig else None

    return {
        "anclas_encontradas": len(anclas),
        "posiciones_anclas": posiciones_anclas,
        "todos_los_anios_match": todos_los_anios,
        "candidatos_cerca_ancla": candidatos,
        "marca_vigente": vigente,
        "hay_ambiguedad": len({c["anio"] for c in candidatos}) > 1,
    }


def _formatear_candidatos(candidatos: list[dict]) -> str:
    """Formatea candidatos como 'año:dist_ancla:dist_vigente' separados por ';'."""
    if not candidatos:
        return ""
    partes = []
    for c in candidatos:
        da = c.get("distancia_ancla", "?")
        dv = c.get("distancia_vigente", "?")
        partes.append(f"{c['anio']}:{da}:{dv}")
    return ";".join(partes)


def _formatear_anios(todos: list[dict]) -> str:
    """Formatea todos los años match a string."""
    if not todos:
        return ""
    return " | ".join(f"{a['anio']}@{a['pos']}" for a in todos)


def _extraer_contexto(texto_original: str, todos_anios: list[dict],
                       candidatos: list[dict], contexto_chars: int = 80) -> str:
    """Extrae un fragmento de texto alrededor de los matches de año."""
    partes = []
    items_a_mostrar = candidatos if candidatos else todos_anios
    for item in items_a_mostrar[:3]:
        pos = item["pos"]
        inicio = max(0, pos - contexto_chars)
        fin = min(len(texto_original), pos + contexto_chars)
        snippet = texto_original[inicio:fin].replace("\n", " ").replace("\r", " ")
        partes.append(f"...{snippet}...")
    return " || ".join(partes)


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print("Uso: python test_cbr_docx.py <ruta_docx>")
        print("Ej:  python test_cbr_docx.py D:\\Diarios\\semana_15_marzo.docx")
        sys.exit(1)

    ruta_docx = sys.argv[1]
    if not os.path.isfile(ruta_docx):
        print(f"ERROR: archivo no encontrado: {ruta_docx}")
        sys.exit(1)

    # ── Configuración de logs persistentes ──
    logs_dir = "logs"
    os.makedirs(logs_dir, exist_ok=True)

    ts_str = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    log_path = os.path.join(logs_dir, f"auditoria_cbr_{ts_str}.log")

    logger = logging.getLogger("auditoria_cbr")
    logger.setLevel(logging.INFO)

    # Evitar duplicación si el logger ya tiene handlers (ej. re-ejecución en notebook)
    if not logger.handlers:
        # Handler de consola
        ch = logging.StreamHandler(sys.stdout)
        ch.setLevel(logging.INFO)
        # Handler de archivo
        fh = logging.FileHandler(log_path, encoding="utf-8")
        fh.setLevel(logging.INFO)
        # Formato común
        formatter = logging.Formatter("%(message)s")
        ch.setFormatter(formatter)
        fh.setFormatter(formatter)
        logger.addHandler(ch)
        logger.addHandler(fh)

    logger.info(f"Log de auditoría: {log_path}")
    logger.info(f"Abriendo DOCX: {ruta_docx}")
    doc = Document(ruta_docx)

    # ── PASO 1: Extraer año del título (misma lógica) ──
    anio_docx = None
    for para in doc.paragraphs:
        texto = para.text.strip()
        if _RE_TITULO_SEMANA.search(texto):
            m_a = re.search(r"(\d{4})", texto)
            if m_a:
                anio_docx = m_a.group(1)
                logger.info(f"  Año del DOCX: {anio_docx}")
            break
    if not anio_docx:
        anio_docx = str(datetime.now().year)
        logger.info(f"  Año no encontrado en título, usando: {anio_docx}")

    # ── PASO 2: Separar párrafos por encabezados de fecha ──
    fecha_actual = None
    bloques_con_fecha: list[tuple[str, str | None]] = []

    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            continue
        if _RE_FECHA_ENCABEZADO.match(texto):
            fecha_actual = f"{texto} {anio_docx}"
            continue
        if texto == "RESUMEN NACIONAL" or _RE_TITULO_SEMANA.search(texto):
            continue
        bloques_con_fecha.append((texto, fecha_actual))

    logger.info(f"  Bloques de causa (párrafos): {len(bloques_con_fecha)}")

    # ── PASO 3: Expandir párrafos con múltiples causas pegadas ──
    bloques_expandidos: list[tuple[str, str | None]] = []
    for texto_bloque, fecha in bloques_con_fecha:
        sub_bloques = separar_bloques(texto_bloque)
        if len(sub_bloques) > 1:
            for sb in sub_bloques:
                bloques_expandidos.append((sb, fecha))
        else:
            bloques_expandidos.append((texto_bloque, fecha))

    logger.info(f"  Bloques después de expandir: {len(bloques_expandidos)}")

    # ── PASO 4: Procesar cada bloque con CBR instrumentado ──
    filas = []
    stats = {
        "total": 0, "excluir": 0, "mantener": 0, "revisar": 0,
        "ambiguos": 0,
        "excluir_ambiguo": [],          # ROLes con exclusión sobre año ambiguo (todos)
        "excluir_ambiguo_v2": [],       # ídem, solo los que pasan pre-filtro v2
        "revisar_sin_anio": 0,          # REVISAR por "año no detectado"
        "revisar_ambiguo": 0,           # REVISAR por "ambiguo"
    }

    for texto_bloque, fecha in bloques_expandidos:
        stats["total"] += 1

        # a. Extraer ROL (regex puro, sin API)
        rol_result = extraer_rol(texto_bloque)
        rol = rol_result[0] if rol_result else ""
        anio_rol = rol_result[1] if rol_result else ""

        # b. Evaluación CBR real
        cbr = evaluar_antiguedad_cbr(texto_bloque)

        # c. Instrumentación de candidatos
        inst = _instrumentar_candidatos(texto_bloque)

        # d. Pre-filtro v2: réplica de PASO 3.5 de parsear_docx_v2
        pasa_prefiltro_v2 = bool(_RE_TIENE_ROL.search(texto_bloque))

        decision = cbr["decision"]
        if decision == "EXCLUIR":
            stats["excluir"] += 1
        elif decision == "MANTENER":
            stats["mantener"] += 1
        else:
            stats["revisar"] += 1

        if inst["hay_ambiguedad"]:
            stats["ambiguos"] += 1

        # Métricas de riesgo
        if inst["hay_ambiguedad"] and decision == "EXCLUIR":
            rol_anio = f"{rol}-{anio_rol}"
            stats["excluir_ambiguo"].append(rol_anio)
            if pasa_prefiltro_v2:
                stats["excluir_ambiguo_v2"].append(rol_anio)

        if decision == "REVISAR":
            motivo = cbr.get("cbr_motivo", "")
            if "no detectado" in motivo.lower():
                stats["revisar_sin_anio"] += 1
            elif "ambiguo" in motivo.lower():
                stats["revisar_ambiguo"] += 1

        todos = inst["todos_los_anios_match"]
        cands = inst["candidatos_cerca_ancla"]
        fragmento = _extraer_contexto(texto_bloque, todos, cands)

        filas.append({
            "ROL": rol,
            "AÑO_ROL": anio_rol,
            "DECISION": decision,
            "PASA_PREFILTRO_V2": "SI" if pasa_prefiltro_v2 else "NO",
            "CBR_AÑO_ELEGIDO": str(cbr["cbr_anio"]) if cbr["cbr_anio"] else "",
            "CBR_MOTIVO": cbr["cbr_motivo"],
            "TODOS_LOS_AÑOS_MATCH": _formatear_anios(todos),
            "CANDIDATOS_CERCA_ANCLA": _formatear_candidatos(cands),
            "TENIA_MARCA_VIGENTE": "SI" if inst["marca_vigente"] else "NO",
            "FRAGMENTO_CONTEXTO": fragmento,
            "BLOQUE_COMPLETO": texto_bloque.replace("\n", "\\n").replace("\r", ""),
        })

    # ── PASO 5: Volcar CSV ──
    fecha_str = datetime.now().strftime("%Y-%m-%d")
    csv_path = f"auditoria_cbr_{fecha_str}.csv"
    columnas = [
        "ROL", "AÑO_ROL", "DECISION", "PASA_PREFILTRO_V2", "CBR_AÑO_ELEGIDO", "CBR_MOTIVO",
        "TODOS_LOS_AÑOS_MATCH", "CANDIDATOS_CERCA_ANCLA", "TENIA_MARCA_VIGENTE",
        "FRAGMENTO_CONTEXTO", "BLOQUE_COMPLETO",
    ]

    with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=columnas, quoting=csv.QUOTE_ALL)
        writer.writeheader()
        writer.writerows(filas)

    logger.info(f"\nCSV generado: {csv_path}")

    # ── PASO 6: Resumen ──
    n_excl_amb = len(stats["excluir_ambiguo"])
    logger.info(f"\n{'=' * 65}")
    logger.info(f"  RESUMEN — Auditoría CBR")
    logger.info(f"{'=' * 65}")
    logger.info(f"  Total bloques procesados : {stats['total']}")
    logger.info(f"  EXCLUIR  (>= {_CBR_ANIO_CORTE}) : {stats['excluir']}")
    logger.info(f"  MANTENER (< {_CBR_ANIO_CORTE})  : {stats['mantener']}")
    logger.info(f"  REVISAR                          : {stats['revisar']}")
    logger.info(f"    └ por \"año no detectado\"      : {stats['revisar_sin_anio']}")
    logger.info(f"    └ por \"ambiguo\"               : {stats['revisar_ambiguo']}")
    logger.info(f"  ───────────────────────────────────────")
    n_excl_amb_v2 = len(stats["excluir_ambiguo_v2"])
    logger.info(f"  EXCLUSIONES SOBRE AÑO AMBIGUO")
    logger.info(f"  (n_candidatos>1 AND EXCLUIR)     : {n_excl_amb}")
    logger.info(f"    ↑ riesgo de falso positivo silencioso")
    if n_excl_amb > 0:
        logger.info(f"    ROLes: {', '.join(stats['excluir_ambiguo'])}")
    logger.info(f"  ───────────────────────────────────────")
    logger.info(f"  EXCLUSIONES SOBRE AÑO AMBIGUO")
    logger.info(f"  (solo PASA_PREFILTRO_V2=True)     : {n_excl_amb_v2}")
    logger.info(f"    ↑ población que producción v2 realmente evalúa")
    if n_excl_amb_v2 > 0:
        logger.info(f"    ROLes: {', '.join(stats['excluir_ambiguo_v2'])}")
    logger.info(f"{'=' * 65}")


if __name__ == "__main__":
    main()