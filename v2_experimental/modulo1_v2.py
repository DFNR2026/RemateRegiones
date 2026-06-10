"""
Módulo 1 v2: Parser optimizado --Regex-Flow + Haiku

Estrategia de reducción de costos:
  1. Pre-filtro regex: bloques sin ROL (C-/V-/A-) se descartan GRATIS
  2. Campos fáciles (ROL, tribunal, demandante, mínimo) → regex puro
  3. Solo dirección/comuna → Claude Haiku (campo genuinamente difícil)
  4. Funciones compartidas (buscar_corte, filtros, etc.) → importadas de v1

Costo estimado: ~$0.05/run vs ~$3.00/run (v1 usa Sonnet para TODO)
"""

import csv
import json
import os
import re
import sys
import logging
from datetime import datetime

import anthropic
import pandas as pd

# Agregar directorio padre al path para importar módulos de producción
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import (
    ANTHROPIC_API_KEY,
    CAUSAS_XLSX,
    SHEET_REFERENCIA, SHEET_CAUSAS,
    DEMANDANTES_EXCLUIDOS, CORTES_RM,
)
from filtro_cbr import evaluar_antiguedad_cbr, _CBR_ANIO_CORTE
from modulo1_parser import (
    # Funciones públicas reutilizables
    limpiar_texto,
    normalizar_ordinal,
    cargar_referencia,
    cargar_historial_roles,
    buscar_corte,
    extraer_rol,
    extraer_tribunal_texto,
    extraer_demandante,
    extraer_direccion_comuna,
    extraer_tipo_propiedad,
    separar_bloques,
)

log = logging.getLogger("m1v2")
log.setLevel(logging.INFO)

_V2_LOGS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs")


def _setup_logging() -> str:
    """Configura logging a archivo + consola. Retorna path del log."""
    os.makedirs(_V2_LOGS_DIR, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_path = os.path.join(_V2_LOGS_DIR, f"v2_{ts}.log")

    # Limpiar handlers previos (evitar duplicados en re-runs)
    log.handlers.clear()

    fmt = logging.Formatter("%(asctime)s [M1v2] %(message)s")

    fh = logging.FileHandler(log_path, encoding="utf-8")
    fh.setLevel(logging.INFO)
    fh.setFormatter(fmt)
    log.addHandler(fh)

    ch = logging.StreamHandler()
    ch.setLevel(logging.INFO)
    ch.setFormatter(fmt)
    log.addHandler(ch)

    return log_path


# ────────────────────────────────────────────────────────────────
# Constantes
# ────────────────────────────────────────────────────────────────

# Pre-filtro: bloques que NO contienen un ROL judicial se descartan sin API
# Superconjunto de extraer_rol: admite puntos de miles y espacios
# en el número (ej. C-2.243-2020), igual que _RE_ROL de modulo1_parser
_RE_TIENE_ROL = re.compile(r'[CcVvAa]\s*[-–—]\s*\d[\d.\s]*[-–—]\s*\d{4}')

# Regex para mínimo de subasta
_RE_MINIMO = re.compile(
    r'[Mm][ií]nimo.*?\$\s*([\d.,]+)|'
    r'(?:suma\s+m[ií]nima|m[ií]nimo\s+(?:de|del)\s+remate).*?\$\s*([\d.,]+)|'
    r'[Mm][ií]nimo.*?(\d+[\d.,]*)\s*(?:UF|U\.F\.)',
    re.IGNORECASE
)

# Copiada de modulo1_parser.py (_normalizar_direccion)
def _normalizar_direccion(direccion: str) -> str:
    """Normaliza dirección: quita prefijos genéricos e indicadores de número."""
    if not direccion:
        return direccion
    direccion = re.sub(r'^(?:calle|avenida)\s+', '', direccion, flags=re.IGNORECASE)
    direccion = re.sub(r'\bN(?:[°º]\.?|um\.)\s*(?=\d)', '', direccion, flags=re.IGNORECASE)
    return direccion.strip()

# Fix 6: Palabras que indican fin del nombre del tribunal
_RE_TRIBUNAL_CORTE = re.compile(
    r'\s+(?:se\s+rematar|rematar[áa]|orden[oó]|se\s+ha\s+ordenado|'
    r'fij[oó]|ubicad[oa]\s+en|en\s+autos|causa\s+rol|en\s+causa|'
    r'con\s+fecha|decret[oó]|dispuso|-\s*ubicad).*$',
    re.IGNORECASE | re.DOTALL
)

# Fix Bug 2: Prefijos que contaminan fuzzy matching
_PREFIJOS_TRIBUNAL = [
    r'^en\s+causa\s+ejecutiva\s+ante\s+(?:el|la)\s+',
    r'^ante\s+(?:el|la)\s+',
    r'^en\s+causa\s+(?:caratulada|seguida)?\s*ante\s+(?:el|la)\s+',
    r'^(?:el|la)\s+',
]

# Fix Bug 3: Ordinales compuestos "Décimo Primer" → 11, etc.
_DECENAS_COMP_V2 = {
    "décimo": 10, "decimo": 10, "décima": 10, "decima": 10,
    "vigésimo": 20, "vigesimo": 20, "vigésima": 20, "vigesima": 20,
    "trigésimo": 30, "trigesimo": 30,
}
_UNIDADES_COMP_V2 = {
    "primer": 1, "primero": 1, "primera": 1,
    "segundo": 2, "segunda": 2,
    "tercer": 3, "tercero": 3, "tercera": 3,
    "cuarto": 4, "cuarta": 4,
    "quinto": 5, "quinta": 5,
    "sexto": 6, "sexta": 6,
    "séptimo": 7, "septimo": 7,
    "octavo": 8, "octava": 8,
    "noveno": 9, "novena": 9,
}


def _normalizar_ordinal_v2(texto: str) -> str:
    """
    Convierte ordinales en palabras a formato numérico con °.
    Fix Bug 3: Maneja compuestos "Décimo Primer" → "11°" (v1 solo maneja vigésimo+).
    """
    from modulo1_parser import ORDINALES
    palabras = texto.split()
    resultado = []
    i = 0
    while i < len(palabras):
        p1 = palabras[i].lower().rstrip(".,;:")
        # Intenta compuesto de dos palabras (décimo+, vigésimo+, trigésimo+)
        if i + 1 < len(palabras):
            p2 = palabras[i + 1].lower().rstrip(".,;:")
            decena = _DECENAS_COMP_V2.get(p1)
            unidad = _UNIDADES_COMP_V2.get(p2)
            if decena is not None and unidad is not None:
                resultado.append(f"{decena + unidad}°")
                i += 2
                continue
        # Simple (reutilizar ORDINALES de v1)
        if p1 in ORDINALES:
            resultado.append(ORDINALES[p1])
            i += 1
            continue
        resultado.append(palabras[i])
        i += 1
    return " ".join(resultado)


# Copiada de modulo1_parser.py (_limpiar_tribunal) + Fix 6 + Bugs 1-3
def _limpiar_tribunal(nombre_raw: str) -> str:
    """Post-regex: limpia nombre de tribunal, trunca en palabras de remate."""
    if not nombre_raw:
        return nombre_raw
    t = nombre_raw

    # ── Bug 1 fix: al splitear por punto, buscar segmento con Juzgado/Tribunal ──
    segmentos = t.split(".")
    if len(segmentos) > 1:
        encontrado = False
        for seg in segmentos:
            if re.search(r'(?i)\b(juzgado|tribunal)\b', seg.strip()):
                t = seg.strip()
                encontrado = True
                break
        if not encontrado:
            # Comportamiento original: primer segmento no-vacío
            for seg in segmentos:
                if seg.strip():
                    t = seg.strip()
                    break

    # ── Bug 2 fix: strip de prefijos que contaminan fuzzy matching ──
    for patron in _PREFIJOS_TRIBUNAL:
        t = re.sub(patron, '', t, flags=re.IGNORECASE).strip()

    # ── Bug 3 fix: normalizar ordinales compuestos (Décimo Primer → 11°) ──
    t = _normalizar_ordinal_v2(t)

    # Fix 14: Normalizar ordinal con ceros: "01°" -> "1°", "02°" -> "2°"
    t = re.sub(r'^0*(\d+)\s*[°º]', lambda m: f"{int(m.group(1))}°", t)
    # Fix 6: Truncar en keywords de remate/legal que no son parte del tribunal
    t = _RE_TRIBUNAL_CORTE.sub('', t)
    # Truncar en coma o punto seguido de espacio (fin natural del nombre)
    t = re.sub(r'[,.](?:\s|$).*$', '', t)
    t = re.sub(r'(\w)-\s+(\w)', r'\1\2', t)
    t = re.sub(r'\s*(?:N[°º]\.?|No\.\s*|Nro\.\s*)\s*\d+.*$', '', t)
    t = re.sub(r'\s+', ' ', t).strip()
    palabras = t.split()
    preposiciones = {'de', 'del', 'la', 'las', 'los', 'el', 'y', 'e', 'en', 'con'}
    resultado = []
    for i, p in enumerate(palabras):
        if i == 0:
            resultado.append(p)
        elif p.lower() in preposiciones:
            resultado.append(p.lower())
        else:
            resultado.append(p)
    return ' '.join(resultado)


# ────────────────────────────────────────────────────────────────
# Claude Haiku -- dirección/comuna + tribunal fallback (Fix 13)
# ────────────────────────────────────────────────────────────────

# Fix 13: Prompt corto para tribunal cuando regex falla
_PROMPT_TRIBUNAL = (
    "Del siguiente texto de aviso de remate judicial chileno, "
    "extrae UNICAMENTE el nombre del tribunal.\n\n"
    "REGLAS:\n"
    "- El tribunal es un Juzgado (Civil, Letras, etc.) de una ciudad.\n"
    "- Incluye el ordinal y la ciudad. Ej: '2° Juzgado Civil de Santiago'.\n"
    "- Si no hay tribunal claro, responde null.\n\n"
    "Responde SOLO con JSON: {{\"tribunal\": \"...\"}}\n\n"
    "Texto: {bloque}"
)


def _extraer_tribunal_haiku(bloque: str) -> str:
    """Llama a Claude Haiku para extraer tribunal. Retorna string o ''."""
    global _api_calls
    try:
        client = _get_claude_client()
        respuesta = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=100,
            messages=[{
                "role": "user",
                "content": _PROMPT_TRIBUNAL.format(bloque=bloque),
            }],
        )
        _api_calls += 1
        texto = respuesta.content[0].text.strip()
        idx = texto.find("{")
        if idx == -1:
            return ""
        campos, _ = json.JSONDecoder().raw_decode(texto, idx)
        if isinstance(campos, dict):
            return str(campos.get("tribunal") or "").strip()
    except Exception as e:
        log.warning(f"Haiku tribunal API error: {e}")
    return ""

_anthropic_client: anthropic.Anthropic | None = None

def _get_claude_client() -> anthropic.Anthropic:
    global _anthropic_client
    if _anthropic_client is None:
        _anthropic_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    return _anthropic_client


# Fix 10: Prompt Haiku pide direccion + comuna explicitamente
_PROMPT_DIRECCION = (
    "Del siguiente texto de aviso de remate judicial chileno, extrae:\n"
    "1. La direccion completa del inmueble que se remata\n"
    "2. La comuna donde se ubica el inmueble\n\n"
    "REGLAS:\n"
    "- La direccion es del INMUEBLE, NO del tribunal.\n"
    "- Busca frases como 'ubicado en', 'propiedad de calle', 'inmueble de'.\n"
    "- La comuna aparece en frases como 'comuna de X', 'ciudad y comuna de X'.\n"
    "- Si la comuna no se menciona explicitamente, inferirla del contexto "
    "(ej: si dice 'Conservador de Bienes Raices de La Serena', la comuna es La Serena).\n"
    "- Si hay parcela, lote, fundo, hijuela, extrae esa descripcion.\n"
    "- Si no hay direccion clara, responde null.\n\n"
    "Responde SOLO con JSON: {{\"direccion\": \"...\", \"comuna\": \"...\"}}\n\n"
    "Texto: {bloque}"
)

# Contadores de llamadas API (para métricas)
_api_calls = 0


# DESACTIVADO [2026-04-29]: direccion es informativa en M1, no justifica costo Haiku.
# Esta funcion y _PROMPT_DIRECCION quedan como referencia pero NO se invocan.
# Para reactivar: restaurar el bloque de fallback en parsear_bloque (buscar el
# log "Haiku dir desactivado"). Ver PROMPT_CC_eliminar_haiku_direccion.md
def _extraer_direccion_haiku(bloque: str) -> tuple[str, str]:
    """
    Llama a Claude Haiku para extraer dirección y comuna.
    Retorna (direccion, comuna). En caso de error retorna ("", "").
    """
    global _api_calls
    try:
        client = _get_claude_client()
        respuesta = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=200,
            messages=[{
                "role": "user",
                "content": _PROMPT_DIRECCION.format(bloque=bloque),
            }],
        )
        _api_calls += 1
        texto = respuesta.content[0].text.strip()
        idx = texto.find("{")
        if idx == -1:
            return "", ""
        campos, _ = json.JSONDecoder().raw_decode(texto, idx)
        if isinstance(campos, dict):
            return (
                str(campos.get("direccion") or "").strip(),
                str(campos.get("comuna") or "").strip(),
            )
    except Exception as e:
        log.warning(f"Haiku API error: {e}")
    return "", ""


# ────────────────────────────────────────────────────────────────
# Extracción regex de mínimo
# ────────────────────────────────────────────────────────────────

def extraer_minimo(bloque: str) -> str:
    """Extrae el monto mínimo del remate ($ o UF)."""
    m = _RE_MINIMO.search(bloque)
    if m:
        for g in m.groups():
            if g:
                return g.strip()
    return ""


# ────────────────────────────────────────────────────────────────
# Parser v2 de un bloque individual
# ────────────────────────────────────────────────────────────────

def _aplanar_texto(texto: str) -> str:
    """Fix 4: Normaliza saltos de línea, tabs y espacios múltiples a espacio simple."""
    texto = re.sub(r'[\n\r\t]+', ' ', texto)
    texto = re.sub(r' {2,}', ' ', texto)
    return texto.strip()


# ── Fix 3: Regex mejorado de demandante con captura proactiva de bancos ──

_RE_BANCOS = re.compile(
    r'\b(Banco\s+(?:de\s+)?Cr[eé]dito\s+[EeYy]\s+Inversiones(?:\s+S\.?A\.?)?|'
    r'Banco\s+Santander(?:\s+Chile)?(?:\s+S\.?A\.?)?|'
    r'Banco\s+(?:de\s+)?(?:Chile|Itaú|Ita[uú]|BICE|Security|'
    r'Falabella|Ripley|Internacional|Consorcio|BBVA|Corpbanca)|'
    r'Banco\s+del?\s+Estado(?:\s+de\s+Chile)?|'
    r'Scotiabank(?:\s+Chile)?|BCI|Corpbanca|BBVA|'
    r'Cooperativa\s+[A-ZÁÉÍÓÚÑ][A-Za-záéíóúñ\s.]+|'
    r'Fondo\s+de\s+Inversi[oó]n\s+[A-ZÁÉÍÓÚÑ][A-Za-záéíóúñ\s.]+|'
    r'Compa[ñn][ií]a\s+de\s+Seguros\s+[A-ZÁÉÍÓÚÑ][A-Za-záéíóúñ\s.]+|'
    r'Financiera\s+[A-ZÁÉÍÓÚÑ][A-Za-záéíóúñ\s.]+|'
    r'Mutual\s+[A-ZÁÉÍÓÚÑ][A-Za-záéíóúñ\s.]+|'
    r'Caja\s+de\s+Compensaci[oó]n\s+[A-ZÁÉÍÓÚÑ][A-Za-záéíóúñ\s.]+)',
    re.IGNORECASE
)

# Fix 9: Patrones de caratulada con todos los separadores chilenos
# Separadores: "con", "CON", "/", "c/", "C/"
_SEP_CARATULA = r'(?:\s+[Cc][Oo][Nn]\s+|/|[Cc]/)'

_RE_CARATULADA_QUOTED = re.compile(
    r'caratulad[ao]s?\s*"([^"]+?)' + _SEP_CARATULA,
    re.IGNORECASE
)

# Fix 9: Caratulada sin comillas — "caratulado BANCO SANTANDER CHILE S.A/GONZALEZ"
_RE_CARATULADA_UNQUOTED = re.compile(
    r'caratulad[ao]s?\s+([A-ZÁÉÍÓÚÑ][^"\n]{3,80}?)' + _SEP_CARATULA + r'[A-ZÁÉÍÓÚÑ]',
    re.IGNORECASE
)

_RE_CARATULADA_CON = re.compile(
    r'(?:caratulad[ao]s?|causa)\s+(?:Rol\s+(?:N[°º]\s*)?C-[\d.]+-\d{4}\s+)?'
    r'"?([A-ZÁÉÍÓÚÑ][^"\n]{3,80}?)' + _SEP_CARATULA + r'[A-ZÁÉÍÓÚÑ]',
)

_RE_DEMANDANTE_SLASH = re.compile(
    r'(?:causa|caratulad[ao]s?)\s+([A-ZÁÉÍÓÚÑ][^/\n]{3,50}?)\s*/\s*[A-ZÁÉÍÓÚÑ]',
    re.IGNORECASE
)


def _extraer_demandante_v2(bloque: str) -> str:
    """
    Extrae demandante con captura proactiva de bancos + patrones de caratula.
    Fix 9: cubre separadores con/CON/ / /c//C/ y caratulada sin comillas.
    """
    # 1. Caratula entre comillas: caratulados "DEMANDANTE con/slash DEMANDADO"
    m = _RE_CARATULADA_QUOTED.search(bloque)
    if m:
        return m.group(1).strip().strip("\"' ")

    # 2. Fix 9: Caratula sin comillas: caratulado BANCO SANTANDER/GONZALEZ
    m = _RE_CARATULADA_UNQUOTED.search(bloque)
    if m:
        candidato = m.group(1).strip().strip("\"' ")
        if len(candidato) > 3 and not re.fullmatch(r"C-[\d.]+-\d{4}", candidato):
            return candidato

    # 3. Caratula con contexto causa: causa [Rol] DEMANDANTE con DEMANDADO
    m = _RE_CARATULADA_CON.search(bloque)
    if m:
        candidato = m.group(1).strip().strip("\"'")
        if len(candidato) > 3 and not re.fullmatch(r"C-[\d.]+-\d{4}", candidato):
            return candidato

    # 4. Formato slash con prefijo causa: causa DEMANDANTE/DEMANDADO
    m = _RE_DEMANDANTE_SLASH.search(bloque)
    if m:
        return m.group(1).strip()

    # 4. Captura proactiva: buscar banco/institucion en contexto de demanda
    m_banco = _RE_BANCOS.search(bloque)
    if m_banco:
        banco = m_banco.group(1).strip()
        pos_banco = m_banco.start()
        # Buscar "con" o "/" después del banco (puede haber espacios intermedios)
        resto = bloque[m_banco.end():]
        if re.match(r'\s*(?:con\s|/)', resto, re.IGNORECASE):
            return banco
        # Buscar "con" en los siguientes 100 chars (el nombre del banco puede
        # estar seguido de "S.A." u otros sufijos antes de "con")
        if re.search(r'\bcon\s+[A-ZÁÉÍÓÚÑ]', resto[:100]):
            return banco
        # Aceptar si el banco está en contexto de caratulada/causa/autos
        contexto_pre = bloque[max(0, pos_banco - 80):pos_banco].lower()
        if any(k in contexto_pre for k in ('caratulad', 'causa', 'autos', 'ejecutiv')):
            return banco

    # 5. Fallback: usar extraer_demandante de v1
    return extraer_demandante(bloque)


# ── Fix 8: Extraccion mejorada de comuna ──

# Comunas conocidas que llevan articulo
_COMUNAS_CON_ARTICULO = {
    "la serena", "la florida", "las condes", "los ángeles", "los angeles",
    "el bosque", "la pintana", "lo barnechea", "la reina", "lo prado",
    "los andes", "la ligua", "la calera", "el monte", "la cisterna",
    "lo espejo", "la granja", "la unión", "la union", "el quisco",
    "las cabras", "la estrella", "el carmen", "lo miranda",
}

_RE_COMUNA_V2 = [
    # "comuna de La Serena" / "comuna de Antofagasta" — capturar articulo si existe
    re.compile(
        r'(?:ciudad\s+y\s+)?comuna\s+(?:y\s+\w+\s+)?de\s+'
        r'((?:La|El|Los|Las|Lo)\s+[A-ZÁÉÍÓÚÑ][A-Za-záéíóúñ]+(?:\s+[A-Za-záéíóúñ]+)?|'
        r'[A-ZÁÉÍÓÚÑ][A-Za-záéíóúñ]+(?:\s+[A-Za-záéíóúñ]+)?)',
    ),
    # "ciudad de X"
    re.compile(
        r'ciudad\s+de\s+'
        r'((?:La|El|Los|Las|Lo)\s+[A-ZÁÉÍÓÚÑ][A-Za-záéíóúñ]+|'
        r'[A-ZÁÉÍÓÚÑ][A-Za-záéíóúñ]+)',
    ),
]

# Terminadores de comuna
_COMUNA_TERMINADORES = re.compile(
    r'\s*(?:,|\.|$|\s+Regi[oó]n|\s+Provincia|\s+inscrit|\s+el\s+dominio)',
    re.IGNORECASE
)


def _extraer_comuna_v2(bloque: str) -> str:
    """
    Extrae comuna con patrones mejorados que incluyen articulos.
    Fix 8: evita cortar "La Serena" a "Serena".
    """
    for pat in _RE_COMUNA_V2:
        for m in pat.finditer(bloque):
            candidato = m.group(1).strip().strip(",. ")
            # Limpiar terminadores
            m_term = _COMUNA_TERMINADORES.search(candidato)
            if m_term and m_term.start() > 0:
                candidato = candidato[:m_term.start()].strip()
            if len(candidato) >= 3:
                return candidato

    # Fallback: buscar "Conservador de Bienes Raices de CIUDAD" (ultimo recurso)
    m = re.search(
        r'Conservador\s+de\s+Bienes\s+Ra[ií]ces\s+de\s+'
        r'((?:La|El|Los|Las|Lo)\s+[A-ZÁÉÍÓÚÑ][A-Za-záéíóúñ]+|'
        r'[A-ZÁÉÍÓÚÑ][A-Za-záéíóúñ]+)',
        bloque, re.IGNORECASE
    )
    if m:
        return m.group(1).strip()

    return ""


# ── Patrones extra de tribunal para v2 (v1 los resuelve con Claude) ──

def _extraer_tribunal_v2(bloque: str) -> str:
    """
    Extrae tribunal con patrones adicionales que v1 delega a Claude.
    Se usa como fallback cuando extraer_tribunal_texto() retorna vacío.
    """
    # Patrón A: "EXTRACTO. ORDINAL JUZGADO [LETRAS|CIVIL|...] CIUDAD rematar|,"
    # Ej: "EXTRACTO. SEGUNDO JUZGADO LETRAS ANTOFAGASTA rematara..."
    # Ej: "EXTRACTO. TERCER JUZGADO LETRAS ANTOFAGASTA, San Martín..."
    m = re.search(
        r'EXTRACTO[.\s]+([A-ZÁÉÍÓÚÑ\d°º]+\s+(?:JUZGADO|Juzgado)[^,\n]{0,60}?)(?:\s+rematar|\s*,)',
        bloque, re.IGNORECASE
    )
    if m:
        return normalizar_ordinal(m.group(1).strip())

    # Patrón B: "EXTRACTO. El Nº Juzgado de Letras de Ciudad -ubicado..."
    # Ej: "EXTRACTO. El Primer Juzgado de Letras de Coquimbo -ubicado en..."
    m = re.search(
        r'EXTRACTO[.\s]+(?:El\s+|La\s+)?(\w+\s+Juzgado[^,\n-]{0,60}?)(?:\s*[-,]|\s+rematar)',
        bloque, re.IGNORECASE
    )
    if m:
        return normalizar_ordinal(m.group(1).strip())

    # Patrón C: "Rol C-XXX-YYYY del Nº Juzgado de ..."
    # Ej: "Rol C-6990-2019 del 2° Juzgado de Letras de Rancagua SE REMATARÁ"
    m = re.search(
        r'[Rr]ol\s+C-[\d.]+-\d{4}\s+del?\s+(\d+[°º]?\s*Juzgado[^,\n]{0,60}?)(?:\s+[Ss][Ee]\s+|,)',
        bloque, re.IGNORECASE
    )
    if m:
        return m.group(1).strip()

    # Patrón D: "causa Rol C-XXX-YYYY, caratulada ..., TRIBUNAL rematar|ordenó"
    # Ej: 'causa Rol C-862-2025, caratulada "SCOTIABANK CHILE CON NARVÁEZ", rematar'
    # El tribunal a veces está antes del Rol y caratulada
    m = re.search(
        r'(?:causa|autos)\s+Rol\s+C-[\d.]+-\d{4}\s*,\s*caratulad[ao]s?\s*"[^"]+"\s*,?\s*'
        r'(\w+\s+Juzgado[^,\n]{0,50}?)(?:\s*,|\s+rematar)',
        bloque, re.IGNORECASE
    )
    if m:
        return normalizar_ordinal(m.group(1).strip())

    return ""


def parsear_bloque_v2(bloque_raw: str, df_ref: pd.DataFrame,
                      historial_roles: set | None = None) -> dict | None:
    """
    Parsea un bloque de remate con estrategia Regex-Flow + Haiku.

    Diferencias vs v1:
      - ROL, tribunal, demandante → regex puro (0 API calls)
      - dirección/comuna → intenta regex primero, Haiku solo si regex falla
      - Mismos filtros y buscar_corte() que v1

    Retorna dict o None.
    """
    bloque = limpiar_texto(bloque_raw)

    # Fix 4: Aplanar texto --reemplazar \n, \r, tabs por espacios simples
    # para que los regex operen sobre texto continuo
    bloque = _aplanar_texto(bloque)

    # Filtrar bloques no-remate
    if "...!!!" in bloque:
        return None
    if re.search(r"juez[a]?\s*(?:[áa]rbitr[oa]|partidor[oa]?)|jue\s*za\s*partidora|partici[oó]n",
                 bloque, re.IGNORECASE):
        return None

    # ── 1. ROL: regex (obligatorio) ──
    rol_result = extraer_rol(bloque)
    if not rol_result:
        return None
    rol, anio = rol_result

    # ── 1b. Historial: filtro temprano ──
    if historial_roles is not None:
        clave = f"{rol}-{anio}"
        if clave in historial_roles:
            return None

    rol_completo = f"C-{rol}-{anio}"

    # ── Audit 11: Filtro de Antigüedad CBR ──
    _cbr = evaluar_antiguedad_cbr(bloque)
    if _cbr["decision"] == "EXCLUIR":
        log.info(f"  [{rol_completo}] Descartada por CBR: dominio inscrito {_cbr['cbr_anio']} >= {_CBR_ANIO_CORTE}")
        return None

    # ── 2. Tribunal: regex puro (v1 strategies + v2 fallback + Haiku) ──
    tribunal_raw = extraer_tribunal_texto(bloque)
    if tribunal_raw:
        log.info(f"  [{rol_completo}] Tribunal regex v1: '{tribunal_raw}'")
    else:
        tribunal_raw = _extraer_tribunal_v2(bloque)
        if tribunal_raw:
            log.info(f"  [{rol_completo}] Tribunal regex v2: '{tribunal_raw}'")
        else:
            # Fix 13: Haiku como ultimo recurso para tribunal
            log.info(f"  [{rol_completo}] Regex no encontro tribunal -- llamando Haiku...")
            tribunal_raw = _extraer_tribunal_haiku(bloque)
            if tribunal_raw:
                log.info(f"  [{rol_completo}] Tribunal Haiku: '{tribunal_raw}'")
            else:
                preview = bloque[:200].replace('\n', ' ')
                log.warning(f"  [{rol_completo}] Haiku no extrajo tribunal. Preview bloque: '{preview}'")
    tribunal_raw = _limpiar_tribunal(tribunal_raw)

    # Filtro partidor/árbitro
    if tribunal_raw:
        t_lower = tribunal_raw.lower()
        if any(term in t_lower for term in ['partidor', 'partidora', 'árbitro', 'árbitra', 'arbitral']):
            log.info(f"  Descartada: tribunal partidor/árbitro: {tribunal_raw}")
            return None

    # ── 3. Demandante: regex mejorado (Fix 3) ──
    demandante = _extraer_demandante_v2(bloque).title()

    # ── 4. Dirección/comuna: solo regex (Haiku dir desactivado 2026-04-29) ──
    direccion, comuna = extraer_direccion_comuna(bloque)

    # Fix 8+11: Siempre correr _extraer_comuna_v2; preferir version con articulo
    comuna_v2 = _extraer_comuna_v2(bloque)
    if comuna_v2:
        if not comuna or len(comuna_v2) > len(comuna):
            comuna = comuna_v2

    # Fix 11: Limpieza post-extraccion de comuna (exceso de captura)
    if comuna:
        for separador in [" sector ", " localidad ", " lugar "]:
            if separador in comuna.lower():
                comuna = comuna.split(separador)[0].strip()

    # Fix 2+5: Validar que la dirección regex sea genuina
    # Direcciones chilenas tienen numeración (dígitos o palabras numéricas)
    if direccion:
        tiene_digito = bool(re.search(r'\d', direccion))
        tiene_numero_texto = bool(re.search(
            r'n[uú]mero|N[°º]|lote|sitio|manzana|departamento|block|piso|mil|ciento',
            direccion, re.IGNORECASE
        ))
        if len(direccion) < 10 or not (tiene_digito or tiene_numero_texto):
            log.info(f"  [{rol_completo}] Direccion regex descartada (corta/sin numero): '{direccion}'")
            direccion = ""

    # DESACTIVADO [2026-04-29]: direccion es informativa en M1 (solo aparece en
    # el DOCX del abogado; no la usan M2/M3/M5 ni el filtrador) y el fallback
    # Haiku generaba ~32% de las llamadas. Ahora si el regex no extrae direccion
    # valida, se deja en blanco. La comuna NO depende de este bloque: la cubren
    # extraer_direccion_comuna (arriba), _extraer_comuna_v2 (arriba) y el Fix 12
    # (extrae del tribunal, mas abajo). Para reactivar: restaurar el bloque que
    # llamaba a _extraer_direccion_haiku. Ver PROMPT_CC_eliminar_haiku_direccion.md
    haiku_used = False
    if not direccion:
        # Diagnostico del regex de direccion (para futura mejora del regex)
        log.debug(f"  [{rol_completo}] Regex no extrajo direccion. Texto (200 chars): '{bloque[:200]}'")
        log.info(f"  [{rol_completo}] Direccion regex vacia -- se deja en blanco (Haiku dir desactivado)")

    direccion = _normalizar_direccion(direccion)

    # ── 5. Tribunal → Corte (RapidFuzz, lookup directo en REFERENCIA) ──
    corte, tribunal_norm = buscar_corte(tribunal_raw, df_ref, rol=rol, anio=anio)

    # Fix 12: Si comuna sigue vacia, extraer ciudad del tribunal normalizado (REFERENCIA)
    if not comuna and tribunal_norm:
        partes = tribunal_norm.split(" de ")
        if len(partes) >= 2:
            comuna = partes[-1].strip()
            log.info(f"  [{rol_completo}] Comuna extraida del tribunal: '{comuna}'")

    # ── 6. Tipo de propiedad ──
    tipo_propiedad = extraer_tipo_propiedad(bloque)

    # ── 7. Filtro RM ──
    region_rm = corte in CORTES_RM
    if region_rm:
        return None

    return {
        "rol": rol,
        "año": anio,
        "corte": corte,
        "tribunal": tribunal_norm or tribunal_raw,
        "tribunal_raw": tribunal_raw,
        "demandante": demandante,
        "demandado": "",  # v2 no extrae demandado (no crítico para pipeline)
        "direccion": direccion,
        "comuna": comuna,
        "tipo_propiedad": tipo_propiedad,
        "region_rm": region_rm,
        "minimo": extraer_minimo(bloque),
        "_haiku_used": haiku_used,
        "_bloque": bloque[:300],
        "cbr_anio": _cbr["cbr_anio"],
        "cbr_flag_revision": _cbr["cbr_flag_revision"],
        "cbr_motivo": _cbr["cbr_motivo"],
    }


# ────────────────────────────────────────────────────────────────
# Función principal: parsear DOCX semanal (v2)
# ────────────────────────────────────────────────────────────────

_MESES_ES = (
    r"(?:ENERO|FEBRERO|MARZO|ABRIL|MAYO|JUNIO|"
    r"JULIO|AGOSTO|SEPT(?:IEMBRE)?|OCTUBRE|NOV(?:IEMBRE)?|DIC(?:IEMBRE)?)"
)
_RE_FECHA_ENCABEZADO = re.compile(rf'^\d{{1,2}}\s+(?:DE\s+)?{_MESES_ES}$')
_RE_TITULO_SEMANA = re.compile(r'REMATES\s+SEMANA', re.IGNORECASE)


_last_log_path: str = ""


def parsear_docx_v2(ruta_docx: str,
                    ignorar_historial: bool = False) -> list[dict]:
    """
    Parsea DOCX semanal con estrategia Regex-Flow + Haiku.
    Retorna lista de dicts con mismo formato que v1 (para comparación).
    """
    global _api_calls, _last_log_path
    _api_calls = 0
    _last_log_path = _setup_logging()

    from docx import Document

    log.info(f"Iniciando Modulo 1 v2 (DOCX) -- archivo: {ruta_docx}")
    log.info(f"Log: {_last_log_path}")

    if not os.path.isfile(ruta_docx):
        log.error(f"Archivo DOCX no encontrado: {ruta_docx}")
        return []

    doc = Document(ruta_docx)

    # Cargar referencia y historial
    df_ref = cargar_referencia()
    log.info(f"REFERENCIA cargada: {len(df_ref)} tribunales")

    if ignorar_historial:
        historial_roles: set = set()
        log.info("Historial: IGNORADO --0 ROLes previos")
    else:
        historial_roles = cargar_historial_roles()
        log.info(f"Historial: {len(historial_roles)} ROLes previos")

    # PASO 1: Extraer año del título
    año_docx = None
    for para in doc.paragraphs:
        texto = para.text.strip()
        if _RE_TITULO_SEMANA.search(texto):
            m_año = re.search(r'(\d{4})', texto)
            if m_año:
                año_docx = m_año.group(1)
            break
    if not año_docx:
        año_docx = str(datetime.now().year)

    # PASO 2: Separar bloques por fecha
    fecha_actual = None
    bloques_con_fecha: list[tuple[str, str | None]] = []

    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            continue
        if _RE_FECHA_ENCABEZADO.match(texto):
            fecha_actual = f"{texto} {año_docx}"
            continue
        if texto == "RESUMEN NACIONAL" or _RE_TITULO_SEMANA.search(texto):
            continue
        bloques_con_fecha.append((texto, fecha_actual))

    log.info(f"Bloques brutos: {len(bloques_con_fecha)}")

    # PASO 3: Expandir párrafos con múltiples causas
    bloques_expandidos: list[tuple[str, str | None]] = []
    for texto_bloque, fecha in bloques_con_fecha:
        sub_bloques = separar_bloques(texto_bloque)
        if len(sub_bloques) > 1:
            for sb in sub_bloques:
                bloques_expandidos.append((sb, fecha))
        else:
            bloques_expandidos.append((texto_bloque, fecha))

    # PASO 3.5: Pre-filtro --descartar bloques sin ROL (GRATIS, sin API)
    bloques_con_rol = []
    bloques_sin_rol_prefiltro = 0
    descartes_prefiltro = []
    for texto_bloque, fecha in bloques_expandidos:
        if _RE_TIENE_ROL.search(texto_bloque):
            bloques_con_rol.append((texto_bloque, fecha))
        else:
            bloques_sin_rol_prefiltro += 1
            descartes_prefiltro.append({
                "fecha_publicacion": fecha or "",
                "texto": texto_bloque[:300],
            })

    log.info(f"Bloques expandidos: {len(bloques_expandidos)} | "
             f"Pre-filtro sin ROL: {bloques_sin_rol_prefiltro} | "
             f"Con ROL: {len(bloques_con_rol)}")

    if descartes_prefiltro:
        _ts = datetime.now().strftime("%Y-%m-%d_%H%M%S")
        _csv_desc = os.path.join("logs", f"descartes_prefiltro_v2_{_ts}.csv")
        os.makedirs("logs", exist_ok=True)
        with open(_csv_desc, "w", newline="", encoding="utf-8-sig") as _f:
            _w = csv.DictWriter(_f, fieldnames=["fecha_publicacion", "texto"],
                                quoting=csv.QUOTE_ALL)
            _w.writeheader()
            _w.writerows(descartes_prefiltro)
        log.warning(
            "PREFILTRO v2 — REVISIÓN MANUAL: %d bloque(s) descartado(s) sin ROL → %s",
            len(descartes_prefiltro),
            _csv_desc,
        )

    # PASO 4: Procesar bloques
    causas_semana: dict[str, dict] = {}
    stats = {
        "bloques_totales": len(bloques_expandidos),
        "prefiltro_sin_rol": bloques_sin_rol_prefiltro,
        "sin_rol": 0,
        "filtrados_año": 0,
        "filtrados_bancoestado": 0,
        "duplicados_semana": 0,
        "sin_tribunal": 0,
        "sin_direccion": 0,
    }

    nombre_docx = os.path.basename(ruta_docx)

    for texto_bloque, fecha_pub in bloques_con_rol:
        causa = parsear_bloque_v2(texto_bloque, df_ref, historial_roles=historial_roles)

        if causa is None:
            stats["sin_rol"] += 1
            continue

        clave = f"{causa['rol']}-{causa['año']}"

        # Filtro año < 2018
        if int(causa["año"]) < 2018:
            stats["filtrados_año"] += 1
            continue

        # Filtro Banco Estado
        dem_lower = causa["demandante"].lower()
        if not dem_lower:
            if re.search(
                r'[Bb]anco\s+(?:del\s+)?[Ee]stado(?:\s+de\s+[Cc]hile)?\s+con\b',
                texto_bloque
            ):
                dem_lower = "banco del estado"
        if any(exc in dem_lower for exc in DEMANDANTES_EXCLUIDOS):
            stats["filtrados_bancoestado"] += 1
            continue

        # Deduplicar
        if clave in causas_semana:
            stats["duplicados_semana"] += 1
            continue

        if not causa["tribunal"]:
            stats["sin_tribunal"] += 1
        if not causa["direccion"]:
            stats["sin_direccion"] += 1

        causa.pop("_bloque", None)
        causa["fecha_publicacion"] = fecha_pub or ""
        causa["fechas_publicacion"] = fecha_pub or ""
        causa["pdf_origen"] = nombre_docx
        causas_semana[clave] = causa

    resultado = list(causas_semana.values())

    # Resumen
    log.info("=" * 60)
    log.info(f"Módulo 1 v2 completado:")
    log.info(f"  Bloques totales:           {stats['bloques_totales']}")
    log.info(f"  Pre-filtro sin ROL:        {stats['prefiltro_sin_rol']}")
    log.info(f"  Sin ROL post-parse:        {stats['sin_rol']}")
    log.info(f"  Filtrados año < 2018:      {stats['filtrados_año']}")
    log.info(f"  Filtrados Banco Estado:    {stats['filtrados_bancoestado']}")
    log.info(f"  Duplicados (semana):       {stats['duplicados_semana']}")
    log.info(f"  Sin tribunal:              {stats['sin_tribunal']}")
    log.info(f"  Sin dirección:             {stats['sin_direccion']}")
    log.info(f"  CAUSAS NUEVAS:             {len(resultado)}")
    log.info(f"  Llamadas API Haiku:        {_api_calls}")
    log.info("=" * 60)

    return resultado


def get_api_calls() -> int:
    """Retorna el número de llamadas API realizadas en esta sesión."""
    return _api_calls


def get_log_path() -> str:
    """Retorna el path del último archivo de log generado."""
    return _last_log_path


def get_logs_dir() -> str:
    """Retorna el directorio de logs de v2."""
    return _V2_LOGS_DIR
