"""
Módulo 1: Parser de PDFs del Diario P&L + Deduplicación + Filtros

Input:  PDFs en D:/Remates/Diarios/ (uno por día)
Output: Lista de dicts con causas nuevas y únicas
"""

import json
import os
import re
import logging
from rapidfuzz import fuzz
from datetime import datetime

import anthropic
import fitz          # PyMuPDF
fitz.TOOLS.mupdf_warnings(False)   # silenciar warnings cosméticos de annotations
import pandas as pd

from config import (
    ANTHROPIC_API_KEY,
    DIARIOS_DIR, CAUSAS_XLSX,
    SHEET_REFERENCIA, SHEET_CAUSAS,
    DEMANDANTES_EXCLUIDOS, CORTES_RM,
)

logging.basicConfig(level=logging.INFO, format="%(asctime)s [M1] %(message)s")
log = logging.getLogger(__name__)

# ────────────────────────────────────────────────────────────────
# Mapeo de ordinales en palabras → número con símbolo °
# ────────────────────────────────────────────────────────────────
ORDINALES = {
    "primer":   "1°",  "primera":   "1°",
    "segundo":  "2°",  "segunda":   "2°",
    "tercer":   "3°",  "tercero":   "3°",  "tercera":  "3°",
    "cuarto":   "4°",  "cuarta":    "4°",
    "quinto":   "5°",  "quinta":    "5°",
    "sexto":    "6°",  "sexta":     "6°",
    "séptimo":  "7°",  "séptima":   "7°",  "septimo":  "7°",  "septima":  "7°",
    "octavo":   "8°",  "octava":    "8°",
    "noveno":   "9°",  "novena":    "9°",
    "décimo":   "10°", "décima":    "10°", "decimo":   "10°", "decima":   "10°",
    "undécimo": "11°", "decimoprimero": "11°",
    "duodécimo": "12°", "decimosegundo": "12°",
    "decimotercero": "13°", "decimocuarto": "14°",
    "decimoquinto":  "15°", "decimosexto":  "16°",
    "decimoséptimo": "17°", "decimoctavo":  "18°",
    "decimonoveno":  "19°",
    "vigésimo":  "20°", "vigesimo":  "20°",
    "vigésima":  "20°", "vigesima":  "20°",
    "trigésimo": "30°",
}

# Compuesto: "vigésimo segundo" → "22°"
DECENAS_COMP = {
    "vigésimo": 20, "vigesimo": 20,
    "trigésimo": 30, "trigesimo": 30,
}
UNIDADES_COMP = {
    "primer": 1, "primero": 1, "primera": 1,
    "segundo": 2, "segunda": 2,
    "tercer": 3, "tercero": 3, "tercera": 3,
    "cuarto": 4, "cuarta": 4,
    "quinto": 5, "quinta": 5,
    "sexto": 6, "sexta": 6,
    "séptimo": 7, "septimo": 7,
    "octavo": 8, "octava": 8,
    "noveno": 9, "novena": 9,
}

# ────────────────────────────────────────────────────────────────
# Utilidades de texto
# ────────────────────────────────────────────────────────────────

def limpiar_texto(texto: str) -> str:
    """Elimina soft-hyphens, normaliza comillas y espacios/saltos de línea."""
    texto = texto.replace("\xad", "")      # soft hyphen (U+00AD)
    texto = texto.replace("\u00ad", "")    # soft hyphen (variante)
    # Normalizar comillas tipográficas → comillas rectas
    texto = texto.replace("\u201c", '"').replace("\u201d", '"')  # " "
    texto = texto.replace("\u2018", "'").replace("\u2019", "'")  # ' '
    texto = texto.replace("\u00ab", '"').replace("\u00bb", '"')  # « »
    texto = re.sub(r"\n+", " ", texto)     # newlines → espacio
    texto = re.sub(r" {2,}", " ", texto)   # múltiples espacios → uno
    # Reconstruir palabras partidas por layout de columnas del PDF
    for palabra_partida, palabra_entera in [
        (r"Juz\s+gado", "Juzgado"),
        (r"JUZGA\s+DO", "JUZGADO"),
        (r"Tri\s+bunal", "Tribunal"),
        (r"BAN\s+CO\b", "BANCO"),
        (r"Con\s+cepci", "Concepci"),
        (r"Valpara\s+[ií]so", "Valparaíso"),
        (r"esta\s+cion", "estacion"),
        (r"Su\s+basta", "Subasta"),
        (r"Se\s+cretaria", "Secretaria"),
        (r"Se\s+cretaría", "Secretaría"),
    ]:
        texto = re.sub(palabra_partida, palabra_entera, texto, flags=re.IGNORECASE)
    return texto.strip()


def normalizar_ordinal(texto: str) -> str:
    """
    Convierte ordinales en palabras a formato numérico con °.
    Ej: "Vigésimo Segundo Juzgado" → "22° Juzgado"
    """
    palabras = texto.split()
    resultado = []
    i = 0
    while i < len(palabras):
        p1 = palabras[i].lower().rstrip(".,;:")
        # Intenta compuesto de dos palabras
        if i + 1 < len(palabras):
            p2 = palabras[i + 1].lower().rstrip(".,;:")
            decena = DECENAS_COMP.get(p1)
            unidad = UNIDADES_COMP.get(p2)
            if decena is not None and unidad is not None:
                resultado.append(f"{decena + unidad}°")
                i += 2
                continue
        # Simple
        if p1 in ORDINALES:
            resultado.append(ORDINALES[p1])
            i += 1
            continue
        resultado.append(palabras[i])
        i += 1
    return " ".join(resultado)


import unicodedata

def _quitar_tildes(texto: str) -> str:
    """Elimina tildes/diacríticos: á→a, é→e, ñ→n, etc."""
    nfkd = unicodedata.normalize("NFKD", texto)
    return "".join(c for c in nfkd if not unicodedata.combining(c))


def _normalizar_para_matching(texto: str) -> str:
    """Normalización agresiva para RapidFuzz."""
    t = texto.lower()
    t = _quitar_tildes(t)
    # Colapsar espacios problemáticos comunes de PDFs
    t = t.replace("val paraiso", "valparaiso")
    t = t.replace("villa rrica", "villarrica")
    t = t.replace("los angeles", "losangeles")
    # Ordinales escritos → numéricos
    t = re.sub(r'\b(primer|primera|1er)\b', '1', t)
    t = re.sub(r'\b(segundo|segunda|2do)\b', '2', t)
    t = re.sub(r'\b(tercer|tercero|tercera|3er)\b', '3', t)
    t = re.sub(r'\b(cuarto|cuarta|4to)\b', '4', t)
    t = re.sub(r'\b(quinto|quinta|5to)\b', '5', t)
    # Abreviaturas estándar
    t = re.sub(r'\bgarantia\b', 'gar', t)
    t = re.sub(r'\bpuerto\b', 'pto', t)
    t = t.replace("jdo", "juzgado")
    # Eliminar símbolo ordinal y puntuación
    t = re.sub(r'[°º\xba\.,-]', ' ', t)
    t = re.sub(r'\bde\b', ' ', t)
    # Solo alfanuméricos y espacios
    t = re.sub(r'[^a-z0-9\s]', ' ', t)
    return re.sub(r'\s+', ' ', t).strip()


# ────────────────────────────────────────────────────────────────
# Carga de datos de referencia
# ────────────────────────────────────────────────────────────────

def cargar_referencia() -> pd.DataFrame:
    """Carga hoja REFERENCIA del Excel (mapeo tribunal → corte)."""
    df = pd.read_excel(CAUSAS_XLSX, sheet_name=SHEET_REFERENCIA)
    df.columns = ["corte", "tribunal"]
    return df


def cargar_historial_roles() -> set:
    """Retorna set de ROLes ya procesados ('XXXXX-YYYY') desde hoja CAUSAS."""
    try:
        df = pd.read_excel(CAUSAS_XLSX, sheet_name=SHEET_CAUSAS)
        if df.empty:
            return set()
        roles = set()
        for _, row in df.iterrows():
            try:
                rol = str(row.get("ROL", "")).strip()
                anio = str(row.get("AÑO", row.get("A\xd1O", ""))).strip()
                if rol and anio:
                    roles.add(f"{rol}-{anio}")
            except Exception:
                continue
        return roles
    except Exception as e:
        log.warning(f"No se pudo leer historial CAUSAS: {e}")
        return set()


# ────────────────────────────────────────────────────────────────
# Limpieza de tribunal (red de seguridad post-LLM)
# ────────────────────────────────────────────────────────────────

def _limpiar_tribunal(nombre_raw: str) -> str:
    """
    Limpia el nombre del tribunal extraído del PDF o de Claude API.
    Se aplica SIEMPRE, como red de seguridad post-LLM.
    """
    if not nombre_raw:
        return nombre_raw

    t = nombre_raw.strip()

    # Regla 4: Unir palabras cortadas por guión de salto de línea
    # "Juzga- do" → "Juzgado", "Corres- pondiente" → "Correspondiente"
    t = re.sub(r'(\w)-\s+(\w)', r'\1\2', t)

    # Regla 5: Eliminar dirección física al final del nombre del tribunal
    # Cortar ante "Nº", "N°", "No.", "Nro" + cualquier cosa
    t = re.sub(r'\s+(Nº|N°|No\.|Nro\.?)\s*\d.*$', '', t, flags=re.IGNORECASE)

    # Cortar ante un número de calle (3+ dígitos al final de la cadena)
    t = re.sub(r'\s+\d{3,}[\s,]?.*$', '', t)

    # Normalizar múltiples espacios
    t = re.sub(r'\s+', ' ', t).strip()

    # Normalizar capitalización: preposiciones en minúscula
    # "Civil De Antofagasta" → "Civil de Antofagasta"
    partes = t.split()
    resultado = []
    for i, p in enumerate(partes):
        if i > 0 and p.lower() in ('de', 'del', 'la', 'las', 'los', 'el'):
            resultado.append(p.lower())
        else:
            resultado.append(p)
    t = ' '.join(resultado)

    return t


_PALABRAS_TRIBUNAL = {
    "juzgado", "civil", "letras", "garantía", "garantia", "mixto", "familia",
    "cobranza", "laboral", "penal", "competencia", "común", "comun", "trabajo",
}

def _extraer_ciudad(nombre_tribunal: str) -> str | None:
    """
    Extrae la ciudad del nombre de un tribunal.
    Ej: '1° Juzgado Civil de Santiago' -> 'santiago'
        'Juzgado de Letras de Rancagua' -> 'rancagua'
        '3° Juzgado de Letras de Iquique' -> 'iquique'  (no 'letras iquique')
    """
    # Buscar la última ocurrencia de "de <Ciudad>" al final
    m = re.search(r'\bde\s+([A-ZÁÉÍÓÚÑ][a-záéíóúñ]+(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+)*)$',
                  nombre_tribunal.strip())
    if m:
        raw = m.group(1)
        # Quitar palabras que son parte del tipo de tribunal, no de la geografía
        palabras = raw.split()
        geo = [p for p in palabras if p.lower() not in _PALABRAS_TRIBUNAL]
        if geo:
            # Normalizar tildes para comparaciones consistentes
            return _quitar_tildes(" ".join(geo).lower())
    return None


def _extraer_ordinal(nombre_tribunal: str) -> str | None:
    """Extrae el número ordinal de un nombre de tribunal (ej: '3° Juzgado...' → '3')."""
    m = re.search(r'(\d+)\s*[°º\xba]', nombre_tribunal)
    if m:
        return m.group(1)
    # Mapeo de ordinales en texto (incluye variantes sin tilde)
    ordinales_texto = {
        'primer': '1', 'primero': '1', 'primera': '1',
        'segundo': '2', 'segunda': '2',
        'tercer': '3', 'tercero': '3', 'tercera': '3',
        'cuarto': '4', 'cuarta': '4',
        'quinto': '5', 'quinta': '5',
        'sexto': '6', 'sexta': '6',
        'séptimo': '7', 'sétimo': '7', 'septimo': '7', 'setimo': '7',
        'octavo': '8', 'octava': '8',
        'noveno': '9', 'novena': '9',
        'décimo': '10', 'decimo': '10',
        'undécimo': '11', 'undecimo': '11',
        'duodécimo': '12', 'duodecimo': '12',
        'decimotercero': '13', 'decimocuarto': '14',
        'decimoquinto': '15', 'decimosexto': '16',
        'decimoséptimo': '17', 'decimoseptimo': '17',
        'decimoctavo': '18', 'decimonoveno': '19',
        'vigésimo': '20', 'vigesimo': '20',
        'trigésimo': '30', 'trigesimo': '30',
    }
    lower = nombre_tribunal.lower()
    # Compuestos: "vigésimo segundo" -> 22
    for dec_txt, dec_val in [('vigésimo', 20), ('vigesimo', 20),
                              ('trigésimo', 30), ('trigesimo', 30)]:
        if dec_txt in lower:
            for uni_txt, uni_val in [('primer', 1), ('primero', 1),
                                      ('segundo', 2), ('tercer', 3), ('tercero', 3),
                                      ('cuarto', 4), ('quinto', 5), ('sexto', 6),
                                      ('séptimo', 7), ('septimo', 7),
                                      ('octavo', 8), ('noveno', 9)]:
                if uni_txt in lower:
                    return str(dec_val + uni_val)
    for texto, num in ordinales_texto.items():
        if texto in lower:
            return num
    return None


# ────────────────────────────────────────────────────────────────
# Matching tribunal → corte
# ────────────────────────────────────────────────────────────────

def buscar_corte(tribunal_raw: str, df_ref: pd.DataFrame,
                 rol: str = "", anio: str = "") -> tuple[str, str]:
    """
    Busca en df_ref la corte que mejor corresponde al tribunal extraído.
    Retorna (corte, tribunal_normalizado).
    Si no encuentra con suficiente confianza, retorna ("DESCONOCIDA", tribunal_raw).
    """
    if not tribunal_raw:
        return "DESCONOCIDA", ""

    etiqueta = f"C-{rol}-{anio}" if rol else ""

    # Normalizar: ordinales, quitar dirección (la dirección suele venir después de coma)
    normalizado = normalizar_ordinal(tribunal_raw)
    # Quitar texto después de la coma (es la dirección del tribunal)
    normalizado = normalizado.split(",")[0].strip()
    # Limpiar ruido habitual
    normalizado = re.sub(r"\s+", " ", normalizado)
    normalizado = normalizado.strip(" .")

    # Post-procesamiento obligatorio: limpiar tribunal (red de seguridad post-LLM)
    normalizado = _limpiar_tribunal(normalizado)

    # Versión normalizada para matching
    norm_match = _normalizar_para_matching(normalizado)

    mejor_score = 0.0
    mejor_corte = "DESCONOCIDA"
    mejor_tribunal = normalizado

    for _, row in df_ref.iterrows():
        trib_ref = str(row["tribunal"])
        ref_match = _normalizar_para_matching(trib_ref)
        # token_set_ratio: ignora orden de palabras y maneja tokens parciales
        score = fuzz.token_set_ratio(norm_match, ref_match)

        if score > mejor_score:
            mejor_score = score
            mejor_corte = str(row["corte"])
            mejor_tribunal = trib_ref

    UMBRAL = 80.0  # RapidFuzz devuelve 0-100; 80 es el punto dulce
    if mejor_score < UMBRAL:
        log.warning(f"  [NO MATCH] {etiqueta} | Tribunal PDF: '{normalizado}' | "
                    f"Mejor Excel: '{mejor_tribunal}' (score: {mejor_score:.1f}%) -> DESCONOCIDA")
        return "DESCONOCIDA", normalizado

    # Validación post-matching: el ordinal del tribunal debe coincidir
    # Evita que "3° de la Serena" matchee a "1° de Letras de la Serena"
    ordinal_pdf = _extraer_ordinal(normalizado)
    ordinal_ref = _extraer_ordinal(mejor_tribunal)
    if ordinal_pdf and ordinal_ref and ordinal_pdf != ordinal_ref:
        log.warning(f"  [ORDINAL MISMATCH] {etiqueta}: PDF dice {ordinal_pdf} pero match es "
                    f"{ordinal_ref} ('{mejor_tribunal}') -- intentando recovery")
        # Recovery: filtrar df_ref solo tribunales con mismo ordinal y re-buscar
        candidatos = []
        for _, row in df_ref.iterrows():
            trib_ref2 = str(row["tribunal"])
            ord_ref2 = _extraer_ordinal(trib_ref2)
            if ord_ref2 == ordinal_pdf:
                candidatos.append((trib_ref2, str(row["corte"])))
        if candidatos:
            mejor_score2 = 0.0
            mejor_corte2 = "DESCONOCIDA"
            mejor_tribunal2 = normalizado
            for trib_ref2, corte_ref2 in candidatos:
                ref_match2 = _normalizar_para_matching(trib_ref2)
                score2 = fuzz.token_set_ratio(norm_match, ref_match2)
                if score2 > mejor_score2:
                    mejor_score2 = score2
                    mejor_corte2 = corte_ref2
                    mejor_tribunal2 = trib_ref2
            if mejor_score2 >= UMBRAL:
                log.info(f"  [ORDINAL RECOVERY] {etiqueta}: re-match '{mejor_tribunal2}' "
                         f"(score: {mejor_score2:.1f}%) | Corte: {mejor_corte2}")
                # Aplicar validación de ciudad al recovery match también
                ciudad_pdf_r = _extraer_ciudad(normalizado)
                ciudad_ref_r = _extraer_ciudad(mejor_tribunal2)
                if ciudad_pdf_r and ciudad_ref_r and ciudad_pdf_r != ciudad_ref_r:
                    score_pen = mejor_score2 * 0.7
                    if score_pen < UMBRAL:
                        log.warning(f"  [CITY MISMATCH post-recovery] {etiqueta}: "
                                    f"'{ciudad_pdf_r}' vs '{ciudad_ref_r}' -- rechazando")
                        return "DESCONOCIDA", normalizado
                    mejor_score2 = score_pen
                return mejor_corte2, mejor_tribunal2
            else:
                log.warning(f"  [ORDINAL RECOVERY FAIL] {etiqueta}: mejor '{mejor_tribunal2}' "
                            f"(score: {mejor_score2:.1f}%) < {UMBRAL}")
        return "DESCONOCIDA", normalizado

    # Validación post-matching: la ciudad del tribunal debe coincidir
    # Evita que "Rancagua" matchee a "Buin" o "Linares" a "Valparaíso"
    ciudad_pdf = _extraer_ciudad(normalizado)
    ciudad_ref = _extraer_ciudad(mejor_tribunal)
    if ciudad_pdf and ciudad_ref and ciudad_pdf != ciudad_ref:
        score_penalizado = mejor_score * 0.7
        if score_penalizado < UMBRAL:
            log.warning(f"  [CITY MISMATCH] {etiqueta}: PDF '{ciudad_pdf}' vs ref '{ciudad_ref}' | "
                        f"score {mejor_score:.1f}% * 0.7 = {score_penalizado:.1f}% < {UMBRAL} -- intentando recovery")
            # Recovery: filtrar df_ref por tribunales que contengan la ciudad del PDF
            candidatos_city = []
            for _, row in df_ref.iterrows():
                trib_ref_c = str(row["tribunal"])
                ciudad_ref_c = _extraer_ciudad(trib_ref_c)
                if ciudad_ref_c and ciudad_ref_c == ciudad_pdf:
                    candidatos_city.append((trib_ref_c, str(row["corte"])))
            if candidatos_city:
                mejor_score_c = 0.0
                mejor_corte_c = "DESCONOCIDA"
                mejor_tribunal_c = normalizado
                for trib_ref_c, corte_ref_c in candidatos_city:
                    ref_match_c = _normalizar_para_matching(trib_ref_c)
                    score_c = fuzz.token_set_ratio(norm_match, ref_match_c)
                    if score_c > mejor_score_c:
                        mejor_score_c = score_c
                        mejor_corte_c = corte_ref_c
                        mejor_tribunal_c = trib_ref_c
                if mejor_score_c >= UMBRAL:
                    # Validar ordinal en el recovery match
                    ord_pdf_c = _extraer_ordinal(normalizado)
                    ord_ref_c = _extraer_ordinal(mejor_tribunal_c)
                    if ord_pdf_c and ord_ref_c and ord_pdf_c != ord_ref_c:
                        log.warning(f"  [CITY RECOVERY ORDINAL FAIL] {etiqueta}: "
                                    f"ordinal {ord_pdf_c} vs {ord_ref_c} -- rechazando")
                        return "DESCONOCIDA", normalizado
                    log.info(f"  [CITY RECOVERY] {etiqueta}: re-match '{mejor_tribunal_c}' "
                             f"(score: {mejor_score_c:.1f}%) | Corte: {mejor_corte_c}")
                    return mejor_corte_c, mejor_tribunal_c
                else:
                    log.warning(f"  [CITY RECOVERY FAIL] {etiqueta}: mejor '{mejor_tribunal_c}' "
                                f"(score: {mejor_score_c:.1f}%) < {UMBRAL}")
            return "DESCONOCIDA", normalizado
        else:
            log.info(f"  [CITY MISMATCH leve] {etiqueta}: PDF '{ciudad_pdf}' vs ref '{ciudad_ref}' | "
                     f"score penalizado {score_penalizado:.1f}% >= {UMBRAL} -- aceptando")
            mejor_score = score_penalizado

    log.info(f"  [MATCH] {etiqueta} | Tribunal: '{normalizado}' -> "
             f"'{mejor_tribunal}' (score: {mejor_score:.1f}%) | Corte: {mejor_corte}")
    return mejor_corte, mejor_tribunal


# ────────────────────────────────────────────────────────────────
# Extracción de campos desde un bloque de texto
# ────────────────────────────────────────────────────────────────

# Patrón ROL judicial: C-XXXXX-YYYY (el número puede tener puntos)
_RE_ROL = re.compile(
    r"[Rr]ol(?:\s*(?:N[°º]|Nro\.?))?\s*C\s*[-–—]?\s*(\d[\d.\s]*\d|\d)\s*[-–—]\s*(\d{4})",
    re.IGNORECASE,
)
# También el formato "EXTRACTO REMATE ROL C-93-2024"
_RE_ROL_HEADER = re.compile(
    r"(?:EXTRACTO\s+)?REMATE\s+ROL\s+C\s*[-–—]?\s*(\d[\d.\s]*\d|\d)\s*[-–—]\s*(\d{4})",
    re.IGNORECASE,
)


def extraer_rol(bloque: str) -> tuple[str, str] | None:
    """Retorna (rol, año) o None si no hay ROL válido."""
    for pat in (_RE_ROL_HEADER, _RE_ROL):
        m = pat.search(bloque)
        if m:
            rol = m.group(1).replace(".", "").replace(" ", "")  # quitar puntos y espacios
            anio = m.group(2)
            return rol, anio
    return None


def extraer_tribunal_texto(bloque: str) -> str:
    """
    Extrae el nombre del tribunal del encabezado del bloque.
    Aplica múltiples estrategias en orden de confianza.
    """
    b = bloque.strip()

    # ── Estrategia 1: "EXTRACTO REMATE ROL C-XXX Ante el Tribunal, ..."
    m = re.match(
        r"EXTRACTO\s+REMATE\s+ROL\s+C-[\d.]+-\d{4}\s+Ante\s+el?\s+(.+?)(?:,|ubicado|se\s+remat)",
        b, re.IGNORECASE
    )
    if m:
        return m.group(1).strip()

    # ── Estrategia 2: "EXTRACTO REMATE: Tribunal, ..."
    m = re.match(
        r"EXTRACTO\s+REMATE[:\s.]+(?:Ante\s+el?\s+)?(.+?)(?:,|\s+ubicado|\s+rematará|\s+rematara|\s+se\s+remat)",
        b, re.IGNORECASE
    )
    if m:
        candidato = m.group(1).strip()
        if _es_tribunal(candidato):
            return candidato

    # ── Estrategia 3: "Remate: Tribunal, ..." o "Remate: Tribunal Rol ..."
    m = re.match(
        r"Remate:\s+(.+?)(?:,|\s+Rol\s+N[°º]|\s+rol\s+C-|\s+se\s+remat)",
        b, re.IGNORECASE
    )
    if m:
        candidato = m.group(1).strip()
        if _es_tribunal(candidato):
            return candidato

    # ── Estrategia 4: "REMATE. Ante [el] Tribunal, ..."
    m = re.match(r"REMATE[. ]+Ante\s+(?:el\s+|la\s+)?(.+?)(?:,|\s+ubicado)", b, re.IGNORECASE)
    if m:
        return m.group(1).strip()

    # ── Estrategia 5: "REMATE JUDICIAL. Tribunal" o "REMATE JUDICIAL TRIBUNAL"
    # Formato corto: "REMATE JUDICIAL. TERCER Juzgado Civil Concepción, ..."
    m = re.match(
        r"REMATE\s+JUDICIAL[.\s]+(?!Causa\s|En\s|\.{2})([A-ZÁÉÍÓÚÑ][\w\s.°º]+?(?:[Jj]uzgado|JUZGADO)[^,\n]{0,60}?)(?:,|\s+[Cc]ausa\s+[Rr]ol|\s+[Ee]n\s+autos|\s+[Oo]rden[oó]|\s+[Ee]n\s+causa)",
        b, re.IGNORECASE
    )
    if m:
        candidato = normalizar_ordinal(m.group(1).strip())
        if _es_tribunal(candidato):
            return candidato

    # ── Estrategia 6: "REMATE JUDICIAL Causa Rol C-XXX «caratulados» Tribunal ordenó"
    # Formato: "REMATE JUDICIAL Causa Rol C-3676-2020, «LAURA/CANALES», Tercer Juzgado Letras Iquique, Ordenó..."
    # (tras limpiar_texto las comillas son comillas rectas normales)
    m = re.search(
        r'"[^"]{3,}?"\s*,?\s*([A-ZÁÉÍÓÚÑ][^"\n,]{5,80}?[Jj]uzgado[^"\n,]{0,50}?)\s*(?:,?\s*[Oo]rden[oó]|\s+\d+\s+de|\s+del\s+d)',
        b, re.IGNORECASE
    )
    if m:
        return normalizar_ordinal(m.group(1).strip())

    # ── Estrategia 7: "REMATE JUDICIAL [ALL CAPS TRIBUNAL] En autos/causa"
    # Formato: "REMATE JUDICIAL PRIMER JUZGADO LA SERENA En autos ejecutivos..."
    m = re.match(
        r"REMATE\s+JUDICIAL\s+([A-Z][A-Z\s]+?(?:JUZGADO|TRIBUNAL)[A-Z\s.°º]*?)\s+(?:[Ee]n\s+autos|[Cc]ausa\s+[Rr]ol|[Ee]n\s+causa)",
        b
    )
    if m:
        return normalizar_ordinal(m.group(1).strip())

    # ── Estrategia 8: "REMATE. 3° Juzgado..." / "REMATE. TERCER JUZGADO..."
    # Incluye "REMATE SEGUNDO. Juzgado..." y "REMATE. 3° Juzgado de Letras..."
    m = re.match(
        r"REMATE[. ]+(\d+[°º\xba]?\s*[Jj]uzgado[^,\n]{0,60}?|[A-ZÁÉÍÓÚÑ][\w\s.°º\xba]*?[Jj]uzgado[^,\n]{0,60}?)"
        r"(?:,|\s+en\s+causa|\s+juicio|\s+[Rr]ol\s+N|\s+se\s+remat|\s+ubicado|\s+[Cc]alle|\s+\d+\s+de|[.]\s)",
        b, re.IGNORECASE
    )
    if m:
        candidato = normalizar_ordinal(m.group(1).strip(" ."))
        if _es_tribunal(candidato):
            return candidato

    # ── Estrategia 9: En el cuerpo "seguida ante el Tribunal"
    m = re.search(r"seguida?\s+ante\s+el?\s+(.+?)(?:,|\s+ubicado)", b, re.IGNORECASE)
    if m:
        return m.group(1).strip()

    # ── Estrategia 10: "en Secretaría/sala del Tribunal" / "en dependencias del Tribunal"
    m = re.search(
        r"en\s+(?:[Ss]ecretar[ií]a\s+del\s+|[Ss]ala\s+del\s+|dependencias\s+del\s+)([^,.\n]{5,80}?[Jj]uzgado[^,.\n]{0,50}?)(?:,|\.|\s+como\s+|\s+el\s+\d|\s+se\s+remat)",
        b, re.IGNORECASE
    )
    if m:
        return normalizar_ordinal(m.group(1).strip())

    # ── Estrategia 11: Buscar "Juzgado" mencionado en el texto del bloque (fallback)
    m = re.search(
        r"(?:del|ante\s+el?|el)\s+((?:\d+[°º]\s+|[A-ZÁÉÍÓÚÑ][a-záéíóúñ]+\s+)?[Jj]uzgado\s+(?:[Cc]ivil|[Dd]e\s+[Ll]etras)[^,\n.]{0,50})",
        b
    )
    if m:
        return normalizar_ordinal(m.group(1).strip())

    return ""


def _es_tribunal(texto: str) -> bool:
    """Verifica que el texto extraído parezca un tribunal (contiene 'juzgado' o 'tribunal')."""
    t = texto.lower()
    return any(k in t for k in ("juzgado", "tribunal", "letras", "civil"))


def extraer_demandante(bloque: str) -> str:
    """
    Extrae el nombre del demandante (la parte antes de 'con' en los caratulados).
    Nota: el bloque ya pasó por limpiar_texto(), comillas normalizadas a "".
    """
    # Patrón: caratulado[s] "DEMANDANTE con DEMANDADO"
    m = re.search(
        r'caratulad[ao]s?\s*"([^"]+?)\s+con\s+[A-ZÁÉÍÓÚÑ]',
        bloque, re.IGNORECASE
    )
    if m:
        return m.group(1).strip().strip("\"' ")

    # Patrón: "en causa DEMANDANTE con DEMANDADO" (sin comillas)
    # NOTA: NO usar re.IGNORECASE aquí — el [A-ZÁÉÍÓÚÑ] final debe ser mayúscula
    # para distinguir "con DEMANDADO" de "con a lo menos", "con fecha", etc.
    m = re.search(
        r'[Cc]ausa\s+(?:[Rr]ol\s+(?:[Nn][°º\xba])?\s*C-[\d.]+-\d{4}\s+)?"?([A-ZÁÉÍÓÚÑ][^"\n]{3,80}?)\s+[Cc][Oo][Nn]\s+[A-ZÁÉÍÓÚÑ]',
        bloque
    )
    if m:
        candidato = m.group(1).strip().strip("\"'")
        if len(candidato) > 3 and not re.fullmatch(r"C-[\d.]+-\d{4}", candidato):
            return candidato

    # Patrón: "Causa DEMANDANTE/DEMANDADO" (slash en vez de con)
    m = re.search(
        r'[Cc]ausa\s+([A-ZÁÉÍÓÚÑ][^/\n]{3,50}?)\s*/\s*[A-ZÁÉÍÓÚÑ]',
        bloque
    )
    if m:
        return m.group(1).strip()

    # Patrón: "BancoXxx/Apellido" directamente (sin "causa")
    m = re.search(
        r'\b([A-ZÁÉÍÓÚÑ][A-Za-záéíóúñüÜ\s.]+?(?:banco|banca|corp|vida|seguro)[A-Za-záéíóúñ\s.]*?)\s*/\s*\w',
        bloque, re.IGNORECASE
    )
    if m:
        return m.group(1).strip()

    return ""


# Patrones de dirección
# Terminadores comunes en escrituras chilenas:
#   "inscrito", "el dominio", "inmueble", "con acceso"  → ya estaban
#   "que corresponde"  → corta tras "calle X número N, que corresponde al sitio..."
#   "resultante de"    → corta tras "calle X, resultante de la subdivisión..."
_RE_DIRECCION = [
    # "ubicado en DIRECCIÓN"
    re.compile(r"ubicad[ao]\s+en\s+(.{10,150}?)(?:\.|,\s*(?:inscrito|el\s+dominio|inmueble|con\s+acceso|que\s+corresponde|resultante\s+de)|$)",
               re.IGNORECASE),
    # "acceso [principal/común] por DIRECCIÓN"
    re.compile(r"acceso\s+(?:principal\s+|com[uú]n\s+)?por\s+(.{10,120}?)(?:,|\.|$)",
               re.IGNORECASE),
    # "propiedad ubicada en DIRECCIÓN"
    re.compile(r"propiedad\s+(?:ubicada?\s+en\s+)(.{10,150}?)(?:\.|,\s*(?:inscrito|el\s+dominio|que\s+corresponde|resultante\s+de)|$)",
               re.IGNORECASE),
    # "inmueble [consistente en] [el ...] ubicado en DIRECCIÓN"
    re.compile(r"inmueble\s+(?:consistente[^,]+,\s+)?(?:[^,]+,\s+)?(?:con\s+acceso\s+(?:por|en)\s+)?(.{5,120}?)(?:,\s*(?:inscrito|el\s+dominio|ciudad)|$)",
               re.IGNORECASE),
]

_RE_COMUNA = re.compile(
    r"(?:ciudad\s+y\s+)?(?:ciudad,\s+)?comuna\s+(?:y\s+\w+\s+)?(?:de\s+la\s+)?(?:de\s+)?([A-ZÁÉÍÓÚÑ][A-Za-záéíóúñ\s]+?)(?:\s*,|\s*\.|$|\s+Regi[oó]n|\s+Provincia)",
    re.IGNORECASE
)

_RE_CALLE_DIRECTA = re.compile(
    r"\b(?:calle|pasaje|avenida|av\.|psje\.)\s+[A-ZÁÉÍÓÚÑ].{5,100}?N[°º]?\s*\d+",
    re.IGNORECASE
)


def extraer_tipo_propiedad(bloque: str) -> str:
    """
    Detecta el tipo de propiedad mencionado en el aviso.

    Reglas (en orden de prioridad):
      "condominio"              → "Ambos"   (unidades pueden ser casa o depto)
      "departamento" / "dpto"   → "Departamentos"
      cualquier otro caso       → "Casas"   (default seguro)
    """
    b = bloque.lower()
    if re.search(r"\bcondominio\b", b):
        return "Ambos"
    if re.search(r"\bdepartamento\b|\bdpto\.?\b", b):
        return "Departamentos"
    return "Casas"


# Palabras clave que indican que el "ubicado en" es la dirección del TRIBUNAL, no del inmueble
_TRIBUNAL_ADDR_CONTEXT = re.compile(
    r"(?:juzgado|tribunal|letras|civil|secretar[ií]a|oficial)\s+(?:\w+\s+){0,5}ubicad",
    re.IGNORECASE
)
# Direcciones conocidas de tribunales (Santiago principalmente)
_TRIBUNAL_ADDR_KNOWN = re.compile(
    r"(?:hu[eé]rfanos|rengifo|barros\s+arana|san\s+mart[ií]n\s+N[°º]?\s*2984|santiago\s+trigo|bilbao\s+777|ciro\s+arredondo|errazuriz\s+s/n)",
    re.IGNORECASE
)


def extraer_direccion_comuna(bloque: str) -> tuple[str, str]:
    """
    Retorna (dirección, comuna) del inmueble.
    Intenta múltiples patrones. Si no encuentra, retorna cadenas vacías.
    Excluye direcciones de tribunales.
    """
    direccion = ""
    comuna = ""

    # Intentar encontrar la fecha de la subasta para buscar la dirección del inmueble
    # DESPUÉS de ese punto (ya que la dirección del tribunal suele aparecer antes)
    m_fecha = re.search(
        r"(?:se\s+remat[aá]r[aá]?|rematar[aá]\s+el\s+d[ií]a|el\s+d[ií]a\s+\d+\s+de\s+\w+\s+de\s+20\d\d)",
        bloque, re.IGNORECASE
    )
    # Buscar desde la fecha si se encontró, si no desde la mitad del bloque
    offset_busqueda = m_fecha.start() if m_fecha else max(0, len(bloque) // 4)

    # Buscar dirección en el texto a partir del offset
    for pat in _RE_DIRECCION:
        for m in pat.finditer(bloque):
            if m.start() < offset_busqueda:
                # Antes del offset, verificar que no sea dirección de tribunal
                contexto_previo = bloque[max(0, m.start()-100):m.start()]
                if _TRIBUNAL_ADDR_CONTEXT.search(contexto_previo):
                    continue
            candidato = m.group(1).strip().strip(",. ")
            # Filtrar candidatos muy cortos
            if len(candidato) < 8:
                continue
            # Filtrar direcciones conocidas de tribunales
            if _TRIBUNAL_ADDR_KNOWN.search(candidato[:60]):
                continue
            # Filtrar si el candidato empieza con términos del tribunal
            if re.match(r"^(?:piso|oficina|sala|secretar)", candidato, re.IGNORECASE):
                continue
            direccion = candidato
            break
        if direccion:
            break

    # Si no encontramos dirección con _RE_DIRECCION, intentar con patrón de calle directa
    if not direccion:
        for m in _RE_CALLE_DIRECTA.finditer(bloque):
            candidato = m.group(0).strip()
            if not _TRIBUNAL_ADDR_KNOWN.search(candidato[:60]):
                direccion = candidato
                break

    # Extraer comuna — buscar todas las ocurrencias y elegir la más relevante
    comunas_encontradas = list(_RE_COMUNA.finditer(bloque))
    for mc in comunas_encontradas:
        candidato_comuna = mc.group(1).strip().strip(",. ")
        # Limpiar residuos
        for ruido in ["inscrito", "el dominio", "provincia", "region", "región"]:
            if ruido in candidato_comuna.lower():
                candidato_comuna = candidato_comuna[:candidato_comuna.lower().find(ruido)].strip(", ")
        if len(candidato_comuna) >= 3:
            # Preferir la primera comuna que aparece después del offset
            if mc.start() >= offset_busqueda or not comuna:
                comuna = candidato_comuna
                if mc.start() >= offset_busqueda:
                    break

    # Limpiar dirección (quitar "commune" al final si quedó pegada)
    if direccion and "comuna" in direccion.lower():
        idx = direccion.lower().find("comuna")
        direccion = direccion[:idx].strip(", ")

    # Truncar dirección muy larga
    if len(direccion) > 150:
        # Cortar en la primera coma si es muy largo
        partes = direccion.split(",")
        direccion = partes[0].strip() if partes else direccion[:150]

    return direccion, comuna


def _normalizar_direccion(direccion: str) -> str:
    """
    Normaliza el campo dirección tras la extracción:
      - Elimina prefijos genéricos "Calle " y "Avenida " (case-insensitive).
        Se mantienen: Pasaje, Condominio, Villa, Población, Conjunto, etc.
      - Elimina indicadores de número "Nº", "N°", "N.°", "Num." antes de un dígito.
        Ej: "San Petersburgo Nº 6351" → "San Petersburgo 6351"
    """
    if not direccion:
        return direccion
    # Eliminar solo "Calle " y "Avenida " al inicio
    direccion = re.sub(r'^(?:calle|avenida)\s+', '', direccion, flags=re.IGNORECASE)
    # Eliminar Nº / N° / N.° / Num. antes de un número
    direccion = re.sub(r'\bN(?:[°º]\.?|um\.)\s*(?=\d)', '', direccion, flags=re.IGNORECASE)
    return direccion.strip()


# ────────────────────────────────────────────────────────────────
# Extracción de campos con Claude API
# ────────────────────────────────────────────────────────────────

_anthropic_client: anthropic.Anthropic | None = None


def _get_claude_client() -> anthropic.Anthropic:
    global _anthropic_client
    if _anthropic_client is None:
        _anthropic_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    return _anthropic_client


_PROMPT_EXTRACCION = (
    "FORMATO OBLIGATORIO: Responde ÚNICAMENTE con un objeto JSON válido. "
    "Sin texto explicativo, sin markdown, sin comentarios. Solo JSON puro. "
    "Si encuentras múltiples remates en el texto, devuelve solo el PRIMERO.\n\n"
    "Eres un extractor de datos de avisos de remates judiciales chilenos. "
    "Del siguiente texto extrae exactamente estos campos en JSON:\n"
    "- rol: solo el número (ej: \"32342\")\n"
    "- año: solo el año (ej: \"2015\")\n"
    "- tribunal: nombre exacto del tribunal\n"
    "- demandante: quien demanda (parte ANTES del \"con\" o \"/\")\n"
    "- demandado: quien es demandado (parte DESPUÉS del \"con\" o \"/\")\n"
    "- direccion: dirección del inmueble rematado\n"
    "- comuna: comuna del inmueble\n"
    "- fecha_remate: fecha en que se realizará el remate en formato DD/MM/YYYY "
    "(ej: \"18/03/2026\"); buscar frases como \"se rematará el día\", \"día NN de mes de YYYY\"\n\n"
    "═══ REGLAS ESTRICTAS (OBLIGATORIAS) ═══\n\n"
    "REGLA 1 — LEE TODO EL TEXTO:\n"
    "- El aviso menciona UN tribunal que ORDENA el remate y UN inmueble que SERÁ rematado.\n"
    "- NO confundas el tribunal que ordena el remate con la comuna donde se ubica el inmueble.\n"
    "- NO confundas el tribunal con el Conservador de Bienes Raíces mencionado en la inscripción.\n"
    "- Lee el texto COMPLETO antes de responder.\n\n"
    "REGLA 2 — EXTRAE LITERALMENTE, NO INTERPRETES:\n"
    "- Si el tribunal dice \"1° Juzgado Civil de Santiago\", escribe exactamente eso, "
    "aunque el inmueble esté en Maipú, Rancagua o cualquier otra comuna.\n"
    "- La comuna es la del INMUEBLE, no la del tribunal.\n"
    "- Tribunal y comuna son campos INDEPENDIENTES.\n\n"
    "REGLA 3 — NUNCA INVENTES ROL:\n"
    "- El ROL tiene formato C-XXXXX-YYYY y aparece explícitamente en el texto.\n"
    "- Si no encuentras un ROL explícito, devuelve rol=null y año=null.\n"
    "- NUNCA deduzcas ni inventes un número de ROL.\n\n"
    "REGLA 4 — RECONSTRUYE PALABRAS CORTADAS:\n"
    "- El PDF a veces corta palabras con guión por salto de línea (ej: \"Juzga-\\ndo\").\n"
    "- Reconstruye la palabra completa: \"Juzga- do\" → \"Juzgado\", \"Antofa- gasta\" → \"Antofagasta\".\n\n"
    "REGLA 5 — EL TRIBUNAL NO INCLUYE SU DIRECCIÓN:\n"
    "- El nombre del tribunal termina en la ciudad (ej: \"1° Juzgado Civil de Antofagasta\").\n"
    "- Si después del nombre aparece una calle o número (ej: \"San Martín Nº 2984\"), "
    "NO lo incluyas en el campo tribunal.\n"
    "- Extrae estrictamente: tipo + ordinal + materia + ciudad.\n\n"
    "REGLA CRÍTICA SOBRE DIRECCIÓN:\n"
    "- La dirección del inmueble DEBE estar textualmente en el aviso.\n"
    "- NUNCA inventes, inferir ni completes una dirección que no aparezca literalmente en el texto.\n"
    "- Si no encuentras una dirección clara del inmueble, devuelve direccion=null.\n"
    "- La dirección puede ser calle con número, pasaje, parcela, lote, sitio, departamento, "
    "block, condominio, edificio, población, villa, hijuela, manzana, fundo, etc.\n"
    "- Busca frases como \"ubicado en\", \"ubicada en\", \"propiedad ubicada en\", "
    "\"inmueble ubicado\", \"inmueble de calle\", \"propiedad de\", \"consistente en\".\n"
    "- Si el aviso describe una parcela o fundo rural, extrae esa descripción completa como dirección.\n"
    "- Si la dirección está fragmentada o repartida en varias líneas, reconstruirla completa.\n"
    "- Si hay texto cortado por formato multi-columna del PDF, unir las partes.\n"
    "- Si no encuentras una dirección explícita pero hay mención de una ubicación o propiedad, "
    "usar esa descripción como dirección.\n"
    "- Prefiere dejar la dirección vacía a inventar algo que no está en el texto.\n\n"
    "Para el campo DEMANDANTE:\n"
    "- Si identificas un nombre de banco o institución financiera (Banco, Financiera, "
    "Cooperativa, Caja, Mutual, etc.), extrae ese nombre como demandante.\n"
    "- Si NO hay banco y el formato es \"APELLIDO/APELLIDO\" o \"APELLIDO con APELLIDO\", "
    "el demandante es el PRIMER apellido/nombre que aparece (el que está ANTES del \"/\" o "
    "\"con\"). En estos casos, agrega el prefijo \"Persona: \" al demandante para "
    "diferenciarlo. Ejemplo: demandante = \"Persona: LAURA\"\n"
    "- En el formato chileno de carátulas judiciales: \"DEMANDANTE con DEMANDADO\" o "
    "\"DEMANDANTE/DEMANDADO\". El primero siempre es el demandante.\n\n"
    "Responde SOLO con el JSON, sin explicaciones. "
    "Si no encuentras un campo responde null.\n\n"
    "Texto: {bloque}"
)


def extraer_campos_claude(bloque: str) -> dict:
    """
    Envía el bloque a Claude Sonnet y retorna los campos extraídos como dict.
    Claves esperadas: rol, año, tribunal, demandante, direccion, comuna.
    En caso de error de API o parseo retorna dict vacío (se activará fallback regex).
    """
    try:
        client = _get_claude_client()
        respuesta = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=256,
            messages=[{
                "role": "user",
                "content": _PROMPT_EXTRACCION.format(bloque=bloque),
            }],
        )
        texto = respuesta.content[0].text.strip()
        # Localizar el primer '{' e intentar parsear solo ese objeto JSON,
        # ignorando cualquier texto que Claude agregue después del cierre '}'
        idx = texto.find("{")
        if idx == -1:
            log.warning("Claude no devolvió JSON válido: %r", texto[:120])
            return {}
        try:
            campos, _ = json.JSONDecoder().raw_decode(texto, idx)
        except json.JSONDecodeError as exc:
            log.warning("JSON inválido en respuesta Claude (%s): %r", exc, texto[:120])
            return {}
        # Si Claude devuelve una lista en vez de dict, tomar el primer elemento
        if isinstance(campos, list):
            campos = campos[0] if campos and isinstance(campos[0], dict) else {}
        if not isinstance(campos, dict):
            return {}
        # Normalizar: convertir None JSON → "" para facilitar comparaciones
        return {k: (v or "") for k, v in campos.items()}
    except Exception as e:
        log.warning(f"Claude API error en extraer_campos_claude: {e}")
        return {}


# ────────────────────────────────────────────────────────────────
# Extracción de un bloque completo
# ────────────────────────────────────────────────────────────────

def parsear_bloque(bloque_raw: str, df_ref: pd.DataFrame,
                   historial_roles: set | None = None) -> dict | None:
    """
    Parsea un bloque de texto de remate y retorna un dict con los campos,
    o None si el bloque no es válido (sin ROL, árbitro, ya en historial, etc.).

    Estrategia de extracción:
      1. ROL:  regex (rápido, confiable; evita llamadas a Claude sin ROL válido).
      1b. Historial: si el ROL ya fue procesado, retorna None SIN llamar a Claude.
      2. tribunal, demandante, direccion, comuna: Claude API (claude-haiku-3-5-20251001).
      3. Fallback a regex cuando Claude devuelve vacío para algún campo.
      4. tribunal → corte: mapeo con hoja REFERENCIA (sin cambios).
    """
    bloque = limpiar_texto(bloque_raw)

    # Filtrar bloques que claramente NO son remates judiciales estándar
    if "...!!!" in bloque:
        return None
    if re.search(r"juez[a]?\s*(?:[áa]rbitr[oa]|partidor[oa]?)|jue\s*za\s*partidora|partici[oó]n", bloque, re.IGNORECASE):
        return None

    # Extraer ROL con regex (requisito obligatorio; si no hay ROL no llamamos a Claude)
    rol_result = extraer_rol(bloque)
    if not rol_result:
        return None
    rol, anio = rol_result

    # ── Filtro temprano: historial (ANTES de llamar a Claude API) ──
    if historial_roles is not None:
        clave = f"{rol}-{anio}"
        if clave in historial_roles:
            return None

    # ── Extracción principal: Claude API ────────────────────────
    campos = extraer_campos_claude(bloque)

    # Logging diagnóstico del bloque raw (para detectar bloques cruzados/corruptos)
    rol_completo = f"C-{rol}-{anio}"

    bloque_preview = bloque[:300].replace('\n', ' ').strip()
    log.debug(f"  Bloque raw [{rol_completo}]: {bloque_preview}")

    def _str_campo(val) -> str:
        """Normaliza un valor de Claude a string: colapsa listas, fuerza str."""
        if isinstance(val, list):
            val = val[0] if val else ""
        if not isinstance(val, str):
            val = str(val) if val else ""
        return val.strip()

    tribunal_raw  = _str_campo(campos.get("tribunal"))
    demandante    = _str_campo(campos.get("demandante")).title()
    demandado     = _str_campo(campos.get("demandado")).title()
    direccion     = _str_campo(campos.get("direccion"))
    comuna        = _str_campo(campos.get("comuna"))
    fecha_remate  = _str_campo(campos.get("fecha_remate"))

    # ── Filtro partidor/árbitro post-extracción (Claude puede extraer tribunal que el regex no atrapó) ──
    if tribunal_raw:
        tribunal_lower = tribunal_raw.lower()
        if any(term in tribunal_lower for term in ['partidor', 'partidora', 'árbitro', 'árbitra', 'arbitral']):
            log.info(f"  Descartada: tribunal es partidor/árbitro: {tribunal_raw}")
            return None

    # ── Retry con prompt focalizado cuando dirección falta ──
    if not direccion or direccion.strip() == '':
        bloque_preview = bloque[:500].replace('\n', '\\n')
        log.warning(f"  [SIN DIRECCIÓN] {rol_completo} — Bloque raw: {bloque_preview}")
        # Segundo intento: prompt focalizado solo en dirección
        log.info(f"  [RETRY DIRECCIÓN] {rol_completo} — Reintentando extracción...")
        try:
            client = _get_claude_client()
            retry_prompt = (
                "Del siguiente texto de un aviso de remate judicial, extrae "
                "ÚNICAMENTE la dirección del inmueble que se remata.\n\n"
                "La dirección puede incluir: calle, número, departamento, block, "
                "condominio, edificio, población, villa, parcela, hijuela, lote, "
                "sitio, manzana, fundo, etc.\n\n"
                "Busca patrones como:\n"
                "- \"ubicado en...\", \"ubicada en...\"\n"
                "- \"inmueble de calle...\", \"propiedad de...\"\n"
                "- \"departamento...\", \"casa...\"\n"
                "- Cualquier mención de calle + número\n"
                "- Descripciones de lotes o sitios rurales\n\n"
                "Responde SOLO con un JSON: {\"direccion\": \"...\", \"comuna\": \"...\"}\n"
                "Si realmente no hay dirección en el texto, responde: "
                "{\"direccion\": null, \"comuna\": null}\n\n"
                f"TEXTO:\n{bloque}"
            )
            response = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=300,
                messages=[{"role": "user", "content": retry_prompt}],
            )
            retry_text = response.content[0].text.strip()
            idx_json = retry_text.find("{")
            if idx_json != -1:
                retry_data, _ = json.JSONDecoder().raw_decode(retry_text, idx_json)
                if isinstance(retry_data, dict) and retry_data.get("direccion"):
                    direccion = str(retry_data["direccion"]).strip()
                    comuna = str(retry_data.get("comuna") or comuna or "").strip()
                    log.info(f"  ✓ Dirección recuperada (retry): {direccion}")
        except Exception as e:
            log.warning(f"  [RETRY DIRECCIÓN] falló: {e}")

    # ── Validación anti-alucinación: dirección debe existir en el texto ──
    if direccion:
        palabras_dir = [p for p in direccion.split() if len(p) > 3][:3]
        if palabras_dir:
            bloque_lower = bloque.lower()
            encontradas = sum(1 for p in palabras_dir if p.lower() in bloque_lower)
            if encontradas < min(2, len(palabras_dir)):
                log.warning(f"  Dirección descartada (no en texto): {direccion}")
                direccion = ""

    # ── Fallback regex para campos que Claude no encontró ───────
    if not tribunal_raw:
        tribunal_raw = extraer_tribunal_texto(bloque)
    if not demandante:
        demandante = extraer_demandante(bloque)
    if not direccion or not comuna:
        dir_re, com_re = extraer_direccion_comuna(bloque)
        if not direccion:
            direccion = dir_re
        if not comuna:
            comuna = com_re

    # Normalizar dirección (quitar prefijos genéricos e indicadores de número)
    direccion = _normalizar_direccion(direccion)

    # Cruzar tribunal con hoja REFERENCIA → corte
    corte, tribunal_norm = buscar_corte(tribunal_raw, df_ref, rol=rol, anio=anio)

    # Determinar tipo de propiedad (búsqueda de palabras clave, no necesita Claude)
    tipo_propiedad = extraer_tipo_propiedad(bloque)

    # NOTA: corte/tribunal → jurisdicción legal; dirección/comuna → ubicación física.
    # Es normal que tribunal esté en Santiago y el inmueble en otra región.
    region_rm = corte in CORTES_RM

    # Filtrar causas de Región Metropolitana — pipeline solo cubre regiones
    if region_rm:
        return None

    return {
        "rol": rol,
        "año": anio,
        "corte": corte,
        "tribunal": tribunal_norm or tribunal_raw,
        "tribunal_raw": tribunal_raw,
        "demandante": demandante,
        "demandado": demandado,
        "fecha_remate": fecha_remate,
        "direccion": direccion,
        "comuna": comuna,
        "tipo_propiedad": tipo_propiedad,
        "region_rm": region_rm,
        "_bloque": bloque[:300],   # para depuración
    }


# ────────────────────────────────────────────────────────────────
# Separación de bloques en el texto del PDF
# ────────────────────────────────────────────────────────────────

_RE_BLOQUE_SEP = re.compile(
    r"(?:EXTRACTO\s+REMATE|REMATE\s+JUDICIAL|REMATE[. ]|Remate:)",
)


def separar_bloques(texto: str) -> list[str]:
    """
    Divide el texto completo del PDF en bloques individuales de remate.
    Cada bloque empieza en un marcador de remate.
    """
    matches = list(_RE_BLOQUE_SEP.finditer(texto))
    bloques = []
    for i, m in enumerate(matches):
        inicio = m.start()
        fin = matches[i + 1].start() if i + 1 < len(matches) else len(texto)
        bloques.append(texto[inicio:fin])
    return bloques


# ────────────────────────────────────────────────────────────────
# Extracción de texto de un PDF
# ────────────────────────────────────────────────────────────────

def extraer_texto_pdf(ruta_pdf: str) -> str:
    """Extrae todo el texto de un PDF usando PyMuPDF."""
    try:
        doc = fitz.open(ruta_pdf)
        partes = []
        for page in doc:
            partes.append(page.get_text("text"))
        doc.close()
        return "\n".join(partes)
    except Exception as e:
        log.error(f"Error extrayendo texto de {ruta_pdf}: {e}")
        return ""


# ────────────────────────────────────────────────────────────────
# Función principal pública
# ────────────────────────────────────────────────────────────────

def _parse_fecha_pdf(nombre: str) -> datetime | None:
    """
    Parsea nombre de PDF con formato DDMMYY a datetime.
    Ej: "010326.pdf" → datetime(2026, 3, 1). Retorna None si el formato no coincide.
    """
    base = os.path.splitext(nombre)[0]
    if len(base) == 6 and base.isdigit():
        try:
            return datetime.strptime(base, "%d%m%y")
        except ValueError:
            pass
    return None


def _formato_fechas(pdfs: list[str]) -> str:
    """
    Convierte lista de nombres de PDF a string legible de fechas de publicación.
    Ej: ["010326.pdf", "020326.pdf", "030326.pdf"] → "01/03 - 03/03/2026"
         ["010326.pdf"]                             → "01/03/2026"
    """
    fechas = sorted(filter(None, (_parse_fecha_pdf(p) for p in pdfs)))
    if not fechas:
        return ""
    if len(fechas) == 1:
        return fechas[0].strftime("%d/%m/%Y")
    return f"{fechas[0].strftime('%d/%m')} - {fechas[-1].strftime('%d/%m/%Y')}"


def parsear_diarios(directorio: str = DIARIOS_DIR,
                    ignorar_historial: bool = False) -> list[dict]:
    """
    Parsea todos los PDFs en el directorio y retorna lista de causas únicas.

    Pasos:
    1. Lee REFERENCIA y CAUSAS del Excel
    2. Para cada PDF extrae y parsea bloques
    3. Deduplica por ROL dentro de la semana
    4. Deduplica contra historial (CAUSAS)
    5. Filtra demandante Banco Estado
    6. Retorna lista final

    Args:
        directorio:        Carpeta con los PDFs del Diario P&L.
        ignorar_historial: Si True, trata el historial CAUSAS como vacío
                           (todas las causas del PDF se consideran nuevas).
                           El Excel NO se modifica. Útil para testing.

    Returns:
        Lista de dicts con campos:
        rol, año, corte, tribunal, demandante, direccion, comuna,
        tipo_propiedad, region_rm
    """
    log.info(f"Iniciando Módulo 1 — directorio: {directorio}")

    # Cargar referencia
    df_ref = cargar_referencia()
    log.info(f"REFERENCIA cargada: {len(df_ref)} tribunales")

    # Cargar historial de ROLes procesados (o ignorarlo en modo testing)
    if ignorar_historial:
        historial_roles: set = set()
        log.info("Historial: IGNORADO (--limpiar-historial activo) — 0 ROLes previos")
    else:
        historial_roles = cargar_historial_roles()
        log.info(f"Historial: {len(historial_roles)} ROLes previos")

    # Listar PDFs
    pdfs = sorted([
        os.path.join(directorio, f)
        for f in os.listdir(directorio)
        if f.lower().endswith(".pdf")
    ])
    if not pdfs:
        log.warning(f"No se encontraron PDFs en {directorio}")
        return []
    log.info(f"PDFs encontrados: {len(pdfs)}: {[os.path.basename(p) for p in pdfs]}")

    # Procesar cada PDF
    causas_semana: dict[str, dict] = {}  # key = "ROL-AÑO"
    stats = {
        "bloques_totales": 0,
        "sin_rol": 0,
        "filtrados_año": 0,
        "filtrados_bancoestado": 0,
        "duplicados_semana": 0,
        "duplicados_historial": 0,
        "sin_tribunal": 0,
        "sin_direccion": 0,
    }

    for ruta_pdf in pdfs:
        nombre = os.path.basename(ruta_pdf)
        log.info(f"Procesando {nombre}...")

        texto = extraer_texto_pdf(ruta_pdf)
        if not texto.strip():
            log.warning(f"  {nombre}: texto vacío")
            continue

        bloques = separar_bloques(texto)
        log.info(f"  {nombre}: {len(bloques)} bloques encontrados")
        stats["bloques_totales"] += len(bloques)

        for bloque in bloques:
            # Pasar historial_roles para filtrar ANTES de llamar a Claude API
            causa = parsear_bloque(bloque, df_ref, historial_roles=historial_roles)

            if causa is None:
                stats["sin_rol"] += 1
                continue

            clave = f"{causa['rol']}-{causa['año']}"

            # Filtro año (descartar causas anteriores a 2018)
            if int(causa["año"]) < 2018:
                log.debug(f"  Filtrado año antiguo: {clave}")
                stats["filtrados_año"] += 1
                continue

            # Filtro Banco Estado
            dem_lower = causa["demandante"].lower()
            if not dem_lower:
                # Fallback: buscar patrón "Banco [del] Estado con" directamente en el bloque
                # (solo cuando está en el rol de demandante explícito)
                if re.search(
                    r'[Bb]anco\s+(?:del\s+)?[Ee]stado(?:\s+de\s+[Cc]hile)?\s+con\b',
                    bloque
                ):
                    dem_lower = "banco del estado"
            if any(exc in dem_lower for exc in DEMANDANTES_EXCLUIDOS):
                log.debug(f"  Filtrado Banco Estado: {clave}")
                stats["filtrados_bancoestado"] += 1
                continue

            # Deduplicar dentro de la semana: acumular PDFs, mantener campos del primero
            if clave in causas_semana:
                stats["duplicados_semana"] += 1
                if nombre not in causas_semana[clave]["_pdfs_origen"]:
                    causas_semana[clave]["_pdfs_origen"].append(nombre)
                continue

            # Advertir si falta tribunal o dirección
            if not causa["tribunal"]:
                stats["sin_tribunal"] += 1
                log.warning(f"  Sin tribunal: {clave} — {causa['_bloque'][:80]}")
            if not causa["direccion"]:
                stats["sin_direccion"] += 1
                log.debug(f"  Sin dirección: {clave}")

            # Limpiar campo interno de depuración antes de guardar
            causa.pop("_bloque", None)
            causa["pdf_origen"]    = nombre
            causa["_pdfs_origen"]  = [nombre]   # lista interna; se convierte abajo
            causas_semana[clave]   = causa

    resultado = list(causas_semana.values())

    # Convertir lista de PDFs a string de fechas legible y eliminar campo interno
    for c in resultado:
        c["fechas_publicacion"] = _formato_fechas(c.pop("_pdfs_origen", []))

    # Resumen
    log.info("=" * 60)
    log.info(f"Módulo 1 completado:")
    log.info(f"  Bloques procesados:        {stats['bloques_totales']}")
    log.info(f"  Sin ROL (descartados):     {stats['sin_rol']}")
    log.info(f"  Filtrados año < 2018:      {stats['filtrados_año']}")
    log.info(f"  Filtrados Banco Estado:    {stats['filtrados_bancoestado']}")
    log.info(f"  Duplicados (semana):       {stats['duplicados_semana']}")
    log.info(f"  Duplicados (historial):    {stats['duplicados_historial']}")
    log.info(f"  Sin tribunal detectado:    {stats['sin_tribunal']}")
    log.info(f"  Sin dirección detectada:   {stats['sin_direccion']}")
    log.info(f"  CAUSAS NUEVAS:             {len(resultado)}")
    log.info("=" * 60)

    return resultado


# ────────────────────────────────────────────────────────────────
# Parser de DOCX semanal consolidado
# ────────────────────────────────────────────────────────────────

# Regex para encabezados de fecha en el DOCX ("16 MARZO", "17 MARZO", etc.)
_MESES_ES = (
    r'(?:ENERO|FEBRERO|MARZO|ABRIL|MAYO|JUNIO|JULIO|AGOSTO|'
    r'SEPTIEMBRE|OCTUBRE|NOVIEMBRE|DICIEMBRE)'
)
_RE_FECHA_ENCABEZADO = re.compile(rf'^\d{{1,2}}\s+(?:DE\s+)?{_MESES_ES}$')
_RE_TITULO_SEMANA = re.compile(r'REMATES\s+SEMANA', re.IGNORECASE)


def parsear_docx_semanal(ruta_docx: str,
                          ignorar_historial: bool = False) -> list[dict]:
    """
    Parsea el DOCX semanal consolidado del Diario P&L.
    Retorna lista de dicts con el mismo formato que parsear_diarios().

    El DOCX tiene estructura:
        Párrafo título: "RESUMEN NACIONAL" / "REMATES SEMANA DEL ... 2026"
        Párrafo fecha:  "16 MARZO"
        Párrafo causa:  texto completo de un aviso de remate
        (vacíos intercalados)

    Reutiliza parsear_bloque(), buscar_corte(), filtros de Banco Estado,
    deduplicación y todos los internals de M1.
    """
    from docx import Document

    log.info(f"Iniciando Módulo 1 (DOCX) — archivo: {ruta_docx}")

    if not os.path.isfile(ruta_docx):
        log.error(f"Archivo DOCX no encontrado: {ruta_docx}")
        return []

    doc = Document(ruta_docx)

    # Cargar referencia y historial (misma lógica que parsear_diarios)
    df_ref = cargar_referencia()
    log.info(f"REFERENCIA cargada: {len(df_ref)} tribunales")

    if ignorar_historial:
        historial_roles: set = set()
        log.info("Historial: IGNORADO — 0 ROLes previos")
    else:
        historial_roles = cargar_historial_roles()
        log.info(f"Historial: {len(historial_roles)} ROLes previos")

    # ── PASO 1: Extraer año del título "REMATES SEMANA DEL ... 2026" ──
    año_docx = None
    for para in doc.paragraphs:
        texto = para.text.strip()
        if _RE_TITULO_SEMANA.search(texto):
            m_año = re.search(r'(\d{4})', texto)
            if m_año:
                año_docx = m_año.group(1)
                log.info(f"Año extraído del título: {año_docx}")
            break

    if not año_docx:
        log.warning("No se encontró año en título del DOCX — usando año actual")
        año_docx = str(datetime.now().year)

    # ── PASO 2: Separar párrafos por encabezados de fecha ──
    fecha_actual = None
    bloques_con_fecha: list[tuple[str, str | None]] = []

    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            continue

        # ¿Es encabezado de fecha? ("16 MARZO")
        if _RE_FECHA_ENCABEZADO.match(texto):
            fecha_actual = f"{texto} {año_docx}"
            log.info(f"  Fecha detectada: {fecha_actual}")
            continue

        # ¿Es título ignorable?
        if texto == "RESUMEN NACIONAL" or _RE_TITULO_SEMANA.search(texto):
            continue

        # Es un bloque de causa
        bloques_con_fecha.append((texto, fecha_actual))

    log.info(f"Bloques de causa encontrados en DOCX: {len(bloques_con_fecha)}")

    # ── PASO 3: Expandir párrafos con múltiples causas pegadas ──
    # (dos ROLes en un mismo párrafo → separar con _RE_BLOQUE_SEP)
    bloques_expandidos: list[tuple[str, str | None]] = []
    for texto_bloque, fecha in bloques_con_fecha:
        sub_bloques = separar_bloques(texto_bloque)
        if len(sub_bloques) > 1:
            log.info(f"  Párrafo con {len(sub_bloques)} causas detectadas — separando")
            for sb in sub_bloques:
                bloques_expandidos.append((sb, fecha))
        else:
            bloques_expandidos.append((texto_bloque, fecha))

    if len(bloques_expandidos) != len(bloques_con_fecha):
        log.info(f"Bloques después de expansión: {len(bloques_expandidos)}")

    # ── PASO 4: Procesar cada bloque (misma lógica que parsear_diarios) ──
    causas_semana: dict[str, dict] = {}
    stats = {
        "bloques_totales": len(bloques_expandidos),
        "sin_rol": 0,
        "filtrados_año": 0,
        "filtrados_bancoestado": 0,
        "duplicados_semana": 0,
        "sin_tribunal": 0,
        "sin_direccion": 0,
    }

    nombre_docx = os.path.basename(ruta_docx)

    for texto_bloque, fecha_pub in bloques_expandidos:
        causa = parsear_bloque(texto_bloque, df_ref, historial_roles=historial_roles)

        if causa is None:
            stats["sin_rol"] += 1
            continue

        clave = f"{causa['rol']}-{causa['año']}"

        # Filtro año < 2018
        if int(causa["año"]) < 2018:
            log.debug(f"  Filtrado año antiguo: {clave}")
            stats["filtrados_año"] += 1
            continue

        # Filtro Banco Estado
        dem_lower = causa["demandante"].lower()
        if not dem_lower:
            if re.search(
                r'[Bb]anco\s+(?:del\s+)?[Ee]stado(?:\s+de\s+[Cc]hile)?\s+con\b',
                texto_bloque
            ):
                dem_lower = "banco del estado"
        if any(exc in dem_lower for exc in DEMANDANTES_EXCLUIDOS):
            log.debug(f"  Filtrado Banco Estado: {clave}")
            stats["filtrados_bancoestado"] += 1
            continue

        # Deduplicar dentro de la semana
        if clave in causas_semana:
            stats["duplicados_semana"] += 1
            continue

        # Advertencias
        if not causa["tribunal"]:
            stats["sin_tribunal"] += 1
            log.warning(f"  Sin tribunal: {clave}")
        if not causa["direccion"]:
            stats["sin_direccion"] += 1
            log.debug(f"  Sin dirección: {clave}")

        # Agregar campos específicos del DOCX
        causa.pop("_bloque", None)
        causa["fecha_publicacion"] = fecha_pub or ""
        causa["fechas_publicacion"] = fecha_pub or ""
        causa["pdf_origen"] = nombre_docx
        causas_semana[clave] = causa

    resultado = list(causas_semana.values())

    # Resumen
    log.info("=" * 60)
    log.info(f"Módulo 1 (DOCX) completado:")
    log.info(f"  Bloques procesados:        {stats['bloques_totales']}")
    log.info(f"  Sin ROL (descartados):     {stats['sin_rol']}")
    log.info(f"  Filtrados año < 2018:      {stats['filtrados_año']}")
    log.info(f"  Filtrados Banco Estado:    {stats['filtrados_bancoestado']}")
    log.info(f"  Duplicados (semana):       {stats['duplicados_semana']}")
    log.info(f"  Sin tribunal detectado:    {stats['sin_tribunal']}")
    log.info(f"  Sin dirección detectada:   {stats['sin_direccion']}")
    log.info(f"  CAUSAS NUEVAS:             {len(resultado)}")
    log.info("=" * 60)

    return resultado


# ────────────────────────────────────────────────────────────────
# Script de prueba standalone
# ────────────────────────────────────────────────────────────────

def _cmd_debug_rol(rol_arg: str) -> None:
    """
    Modo diagnóstico: busca un ROL específico en los PDFs e imprime:
      1. Nombre del PDF donde se encontró el bloque
      2. Texto raw completo del bloque (sin ningún procesamiento)
      3. Valores extraídos: demandante, tribunal, direccion, comuna
      4. Resultado de cada regex de dirección aplicado al texto raw
    """
    # Normalizar argumento: "C-32342-2015" o "32342-2015"
    m = re.match(r"^[Cc]-?(\d[\d.]*)-(\d{4})$", rol_arg.strip())
    if not m:
        print(f"ERROR: formato de ROL inválido: {rol_arg!r}")
        print("  Esperado: C-32342-2015")
        return

    rol_buscado = m.group(1).replace(".", "")
    año_buscado = m.group(2)
    clave_buscada = f"{rol_buscado}-{año_buscado}"

    print(f"Buscando ROL: C-{rol_buscado}-{año_buscado}")
    print(f"Directorio: {DIARIOS_DIR}\n")

    df_ref = cargar_referencia()

    pdfs = sorted([
        os.path.join(DIARIOS_DIR, f)
        for f in os.listdir(DIARIOS_DIR)
        if f.lower().endswith(".pdf")
    ])
    if not pdfs:
        print(f"No se encontraron PDFs en {DIARIOS_DIR}")
        return

    encontrado = False
    for ruta_pdf in pdfs:
        nombre_pdf = os.path.basename(ruta_pdf)
        texto = extraer_texto_pdf(ruta_pdf)
        if not texto.strip():
            continue

        bloques = separar_bloques(texto)
        for i, bloque_raw in enumerate(bloques):
            # Buscar el ROL en el texto limpio (como hace el parser normal)
            bloque_limpio = limpiar_texto(bloque_raw)
            rol_result = extraer_rol(bloque_limpio)
            if not rol_result:
                continue
            rol, año = rol_result
            if f"{rol}-{año}" != clave_buscada:
                continue

            encontrado = True
            sep = "=" * 70

            # 1. PDF
            print(sep)
            print(f"1. PDF: {nombre_pdf}  (bloque #{i + 1} de {len(bloques)})")
            print(sep)

            # 2. Texto raw sin ningún procesamiento
            print(f"\n2. TEXTO RAW DEL BLOQUE:")
            print("-" * 70)
            print(bloque_raw)
            print("-" * 70)

            # 3. Valores extraídos
            print(f"\n3. VALORES EXTRAIDOS:")
            causa = parsear_bloque(bloque_raw, df_ref)
            if causa:
                print(f"   demandante : {causa['demandante']!r}")
                print(f"   tribunal   : {causa['tribunal']!r}")
                print(f"   direccion  : {causa['direccion']!r}")
                print(f"   comuna     : {causa['comuna']!r}")
            else:
                print("   (parsear_bloque retorno None -- bloque invalido/filtrado)")

            # 4. Regex de dirección sobre texto raw
            print(f"\n4. RESULTADO DE _RE_DIRECCION SOBRE TEXTO RAW:")
            print("-" * 70)
            nombres_pat = [
                'ubicad[ao] en ...',
                'acceso [principal] por ...',
                'propiedad ubicada en ...',
                'inmueble ... ubicado en ...',
            ]
            for j, pat in enumerate(_RE_DIRECCION):
                etiqueta = nombres_pat[j] if j < len(nombres_pat) else f"pat[{j}]"
                matches = list(pat.finditer(bloque_raw))
                if matches:
                    for k, mx in enumerate(matches):
                        print(f"   [{j}] {etiqueta}  =>  match #{k + 1}")
                        print(f"        span    : {mx.span()}")
                        print(f"        full    : {mx.group(0)!r}")
                        print(f"        group(1): {mx.group(1)!r}")
                else:
                    print(f"   [{j}] {etiqueta}  =>  None")
            print("-" * 70)
            print()

    if not encontrado:
        print(f"ROL C-{rol_buscado}-{año_buscado} no encontrado en ningún PDF.")


if __name__ == "__main__":
    import sys

    if "--debug-rol" in sys.argv:
        idx = sys.argv.index("--debug-rol")
        if idx + 1 >= len(sys.argv):
            print("ERROR: --debug-rol requiere un argumento. Ejemplo: --debug-rol C-32342-2015")
            sys.exit(1)
        _cmd_debug_rol(sys.argv[idx + 1])
        sys.exit(0)
    else:
        causas = parsear_diarios()
        print(f"\n{'='*60}")
        print(f"CAUSAS ENCONTRADAS: {len(causas)}")
        print(f"{'='*60}")
        for c in causas:
            region = "RM" if c["region_rm"] else "REGION"
            print(f"[{region}] C-{c['rol']}-{c['año']} | {c['tribunal']} | {c['demandante']}")
            print(f"       Direccion  : {c['direccion']}")
            print(f"       Comuna     : {c['comuna']}")
            print(f"       Tipo prop. : {c['tipo_propiedad']}")
            print(f"       Corte      : {c['corte']}")
            print()
